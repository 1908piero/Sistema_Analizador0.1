from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import pandas as pd
import io


APA_FONT = "Times New Roman"
APA_SIZE = Pt(12)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    if top is not None:
        tcBorders.append(parse_xml(
            f'<w:top {nsdecls("w")} w:val="single" w:sz="{top}" w:space="0" w:color="000000"/>'))
    if bottom is not None:
        tcBorders.append(parse_xml(
            f'<w:bottom {nsdecls("w")} w:val="single" w:sz="{bottom}" w:space="0" w:color="000000"/>'))
    if left is not None:
        tcBorders.append(parse_xml(
            f'<w:start {nsdecls("w")} w:val="{left}"/>'))
    if right is not None:
        tcBorders.append(parse_xml(
            f'<w:end {nsdecls("w")} w:val="{right}"/>'))
    tcPr.append(tcBorders)


def _remove_vertical_borders(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, left="nil", right="nil")


def _set_row_borders(row, sz_top=None, sz_bottom=None):
    for cell in row.cells:
        _set_cell_borders(cell, top=sz_top, bottom=sz_bottom, left="nil", right="nil")


class APA7Exporter:
    def __init__(self):
        self.doc = Document()
        self._setup_document()
        self.table_counter = 0
        self.figure_counter = 0

    def _setup_document(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = APA_FONT
        font.size = APA_SIZE
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    def _add_title(self, text: str, level: int = 1):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.name = APA_FONT
        run.font.size = Pt(14 if level == 0 else 12)
        p.paragraph_format.space_after = Pt(12)
        return p

    def _add_heading_apa(self, text: str, level: int = 1):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = APA_FONT
        run.font.size = APA_SIZE
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def _add_paragraph(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = APA_FONT
        run.font.size = APA_SIZE
        return p

    def _format_table_title(self, title: str):
        self.table_counter += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"Tabla {self.table_counter}")
        run.bold = True
        run.font.name = APA_FONT
        run.font.size = APA_SIZE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(12)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run(title)
        run2.italic = True
        run2.font.name = APA_FONT
        run2.font.size = APA_SIZE
        p2.paragraph_format.space_after = Pt(6)
        return p2

    def _format_figure_title(self, title: str):
        self.figure_counter += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"Figura {self.figure_counter}")
        run.bold = True
        run.font.name = APA_FONT
        run.font.size = APA_SIZE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(12)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run(title)
        run2.italic = True
        run2.font.name = APA_FONT
        run2.font.size = APA_SIZE
        p2.paragraph_format.space_after = Pt(6)
        return p2

    def _add_note(self, text: str = "Datos procesados automáticamente por el software estadístico."):
        p = self.doc.add_paragraph()
        run_label = p.add_run("Nota. ")
        run_label.italic = True
        run_label.font.name = APA_FONT
        run_label.font.size = Pt(10)
        run_text = p.add_run(text)
        run_text.font.name = APA_FONT
        run_text.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        return p

    def _apa_table_from_df(self, df: pd.DataFrame, title: str):
        self._format_table_title(title)
        n_rows, n_cols = df.shape
        table = self.doc.add_table(rows=n_rows + 1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(col_name))
            run.bold = True
            run.font.name = APA_FONT
            run.font.size = Pt(10)

        for i in range(n_rows):
            for j in range(n_cols):
                cell = table.rows[i + 1].cells[j]
                val = df.iloc[i, j]
                if isinstance(val, float):
                    val = f"{val:.2f}"
                else:
                    val = str(val)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.font.name = APA_FONT
                run.font.size = Pt(10)

        _remove_vertical_borders(table)
        _set_row_borders(table.rows[0], sz_top="12", sz_bottom="6")
        for i in range(1, n_rows):
            _set_row_borders(table.rows[i])
        _set_row_borders(table.rows[n_rows], sz_bottom="12")
        self._add_note("Elaboración propia a partir de los datos procesados.")

    def export_sampling(self, result: dict):
        self._add_title("Cálculo del Tamaño de Muestra - Muestreo Aleatorio Simple (M.A.S.)", level=0)
        data = [
            ["Nivel de confianza", f"{result['confidence']}%"],
            ["Valor Z", f"{result['Z']:.3f}"],
            ["Probabilidad de éxito (p)", f"{result['p']}"],
            ["Probabilidad de fracaso (q)", f"{result['q']:.2f}"],
            ["Error admisible (e)", f"{result['e']}"],
        ]
        if result["N"] is not None:
            data.append(["Tamaño de la población (N)", str(result["N"])])
            data.append(["Muestra inicial (n₀)", str(result["n0"])])
            data.append(["Muestra corregida (nf)", str(result["nf"])])
        else:
            data.append(["Tamaño de muestra (n)", str(result["n_unknown"])])
        df = pd.DataFrame(data, columns=["Parámetro", "Valor"])
        self._apa_table_from_df(df, "Parámetros de entrada y resultados del cálculo muestral")

    def export_variable_classification(self, classification: dict):
        data = [["Variable", "Tipo"]]
        for var, tipo in classification.items():
            data.append([var, tipo.replace("_", " ").title()])
        df = pd.DataFrame(data[1:], columns=data[0])
        self._apa_table_from_df(df, "Clasificación automática de variables del dataset")

    def export_frequency_table(self, freq_result: dict, measures: dict = None):
        var_name = freq_result["var_name"]
        var_type = freq_result["var_type"]
        table = freq_result["table"]
        is_grouped = freq_result["is_grouped"]
        type_label = var_type.replace("_", " ").title()
        title = f"Distribución de frecuencias de {var_name} ({type_label})"
        if is_grouped:
            R = freq_result.get("R", 0)
            m = freq_result.get("m", 0)
            C = freq_result.get("C", 0)
            self._add_paragraph(
                f"Rango (R) = {R:.2f}  |  Intervalos (Sturges: m = 1 + 3.322·log(n)) = {m}  |  Amplitud (C) = {C:.2f}"
            )
            display_cols = ["Intervalo", "Xi", "fi", "Fi", "hi", "hi%", "Hi%"]
        else:
            display_cols = [table.columns[0], "fi", "Fi", "hi", "hi%", "Hi%"]
        display_df = table[display_cols].copy()
        for col in display_df.columns:
            if display_df[col].dtype in ["float64", "float32"]:
                display_df[col] = display_df[col].apply(lambda x: f"{x:.2f}")
        self._apa_table_from_df(display_df, title)

    def export_measures(self, measures: dict, var_name: str):
        if measures is None:
            return
        self._add_heading_apa(f"Medidas Estadísticas - {var_name}", level=2)
        if measures.get("type") == "cualitativa":
            data = [
                ["Tamaño de muestra (n)", str(measures["n"])],
                ["Moda (Mo)", str(measures["mode"])],
            ]
            df = pd.DataFrame(data, columns=["Medida", "Valor"])
            self._apa_table_from_df(df, "Medidas para variable cualitativa")
            return
        central_data = [
            ["Media aritmética (X̄)", f"{measures['mean']:.2f}"],
            ["Mediana (Me)", f"{measures['median']:.2f}"],
            ["Moda (Mo)", f"{measures['mode']:.2f}"],
            ["Media geométrica (X̄g)", f"{measures['geometric_mean']:.2f}"],
            ["Media armónica (Mh)", f"{measures['harmonic_mean']:.2f}"],
        ]
        df_central = pd.DataFrame(central_data, columns=["Medida", "Valor"])
        self._apa_table_from_df(df_central, "Medidas de tendencia central")
        dispersion_data = [
            ["Rango", f"{measures['range']:.2f}"],
            ["Varianza muestral (S²)", f"{measures['variance']:.2f}"],
            ["Desviación estándar (S)", f"{measures['std_dev']:.2f}"],
            ["Coeficiente de variación (CV%)", f"{measures['cv']:.2f}%"],
        ]
        df_disp = pd.DataFrame(dispersion_data, columns=["Medida", "Valor"])
        self._apa_table_from_df(df_disp, "Medidas de dispersión")
        position_data = [
            ["Cuartil 1 (Q₁)", f"{measures['Q1']:.2f}"],
            ["Cuartil 2 (Q₂) = Mediana", f"{measures['Q2']:.2f}"],
            ["Cuartil 3 (Q₃)", f"{measures['Q3']:.2f}"],
            ["Decil 1 (D₁)", f"{measures['D1']:.2f}"],
            ["Decil 5 (D₅)", f"{measures['D5']:.2f}"],
            ["Decil 9 (D₉)", f"{measures['D9']:.2f}"],
            ["Percentil 10 (P₁₀)", f"{measures['P10']:.2f}"],
            ["Percentil 25 (P₂₅)", f"{measures['P25']:.2f}"],
            ["Percentil 50 (P₅₀)", f"{measures['P50']:.2f}"],
            ["Percentil 75 (P₇₅)", f"{measures['P75']:.2f}"],
            ["Percentil 90 (P₉₀)", f"{measures['P90']:.2f}"],
        ]
        df_pos = pd.DataFrame(position_data, columns=["Medida", "Valor"])
        self._apa_table_from_df(df_pos, "Medidas de posición (cuartiles, deciles y percentiles)")
        shape_data = [
            ["Coeficiente de asimetría (Sesgo)", f"{measures['skewness']:.2f}"],
            ["Coeficiente de curtosis (Exceso)", f"{measures['kurtosis']:.2f}"],
        ]
        df_shape = pd.DataFrame(shape_data, columns=["Medida", "Valor"])
        self._apa_table_from_df(df_shape, "Medidas de forma (asimetría y curtosis)")

    def export_chart(self, chart_bytes: io.BytesIO, title: str):
        self._format_figure_title(title)
        self.doc.add_picture(chart_bytes, width=Inches(5.5))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, path: str):
        self.doc.save(path)

    def export_full_analysis(self, df: pd.DataFrame, classification: dict,
                             sampling_result: dict = None,
                             filepath: str = "Reporte_Estadistico_APA7.docx"):
        from model.statistics import FrequencyAnalyzer, MeasuresCalculator, DatasetSummary
        from model.charts import ChartGenerator

        if sampling_result:
            self.export_sampling(sampling_result)
            self.doc.add_page_break()

        self._add_title("Análisis del Dataset", level=0)
        self.export_variable_classification(classification)

        for col in df.columns:
            var_type = classification.get(col, None)
            if var_type is None or var_type == "desconocido":
                continue

            freq_result = FrequencyAnalyzer.compute(df[col], var_type, col)
            if freq_result is None:
                continue
            measures = MeasuresCalculator.compute(freq_result)
            charts = ChartGenerator.generate_all_charts(freq_result, col)

            self._add_heading_apa(f"Análisis detallado de: {col}", level=1)

            self.export_frequency_table(freq_result, measures)

            if measures.get("type") == "cualitativa":
                mode_val = measures.get("mode", "N/A")
                fi = 0
                pct = 0.0
                table = freq_result.get("table")
                if table is not None and not table.empty:
                    mode_row = table[table.iloc[:, 0].astype(str) == str(mode_val)]
                    if not mode_row.empty:
                        fi = int(mode_row["fi"].values[0])
                        pct = float(mode_row["hi%"].values[0])
                texto_interp = (
                    f"Al observar la variable «{col}», se encontró que la "
                    f"categoría que más se repite es «{mode_val}», con un total "
                    f"de {fi} casos, lo que equivale al {pct:.2f}% del "
                    f"conjunto de datos analizados."
                )
            else:
                mean_v = measures["mean"]
                std_v = measures["std_dev"]
                med_v = measures["median"]
                q3_v = measures.get("Q3", 0)
                texto_interp = (
                    f"Para la variable «{col}», el valor promedio obtenido es de "
                    f"{mean_v:.2f}, con una dispersión típica de {std_v:.2f} "
                    f"unidades alrededor de la media. La mitad de las "
                    f"observaciones se sitúa por debajo de {med_v:.2f}, mientras "
                    f"que tres cuartas partes de los datos no superan los "
                    f"{q3_v:.2f}."
                )

            self._add_heading_apa("Interpretación", level=2)
            p = self.doc.add_paragraph()
            run = p.add_run(texto_interp)
            run.font.name = APA_FONT
            run.font.size = APA_SIZE
            p.paragraph_format.space_after = Pt(8)

            if measures:
                self.export_measures(measures, col)

            if charts:
                self._add_heading_apa("Gráficos Estadísticos", level=2)
                for chart_key, chart_bytes in charts.items():
                    titles = {
                        "bar": f"Gráfico de barras de {col}",
                        "pie": f"Gráfico de sectores de {col}",
                        "bar_ogive": f"Gráfico de barras con ojiva de {col}",
                        "histogram": f"Histograma de frecuencias de {col}",
                        "freq_poly_ogive": f"Polígono de frecuencias y ojiva de {col}",
                    }
                    self.export_chart(chart_bytes, titles.get(chart_key, f"Gráfico de {col}"))
                    self._add_note()

            self.doc.add_page_break()

        self.save(filepath)
        return filepath
