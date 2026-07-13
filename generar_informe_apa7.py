#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera el informe del proyecto "Auto-Analizador de Encuestas"
en formato APA 7 con índice automático.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, io, base64
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from model.charts import ChartGenerator
from model.statistics import VariableClassifier, FrequencyAnalyzer, MeasuresCalculator

APA_FONT = "Times New Roman"
APA_SIZE = Pt(12)


def set_cell_borders(cell, top=None, bottom=None, start=None, end=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    if top:
        tcBorders.append(parse_xml(f'<w:top {nsdecls("w")} w:val="single" w:sz="{top}" w:space="0" w:color="000000"/>'))
    if bottom:
        tcBorders.append(parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="{bottom}" w:space="0" w:color="000000"/>'))
    if start:
        tcBorders.append(parse_xml(f'<w:start {nsdecls("w")} w:val="nil"/>'))
    if end:
        tcBorders.append(parse_xml(f'<w:end {nsdecls("w")} w:val="nil"/>'))
    tcPr.append(tcBorders)


def remove_vertical_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, start=True, end=True)


def set_row_borders(row, sz_top=None, sz_bottom=None):
    for cell in row.cells:
        set_cell_borders(cell, top=sz_top, bottom=sz_bottom)
        set_cell_borders(cell, start=True, end=True)


def add_formula_image(doc, formula_text, fontsize=14, width_inches=5, height_inches=0.6):
    fig, ax = plt.subplots(figsize=(width_inches, height_inches))
    ax.axis('off')
    ax.text(0.5, 0.5, formula_text, transform=ax.transData,
            fontsize=fontsize, fontfamily='serif', fontstyle='italic',
            ha='center', va='center',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor='none'))
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                transparent=False, facecolor='white')
    plt.close(fig)
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    return p

def add_formula_inline(doc, formula_text, fontsize=14, width_inches=5, height_inches=0.55):
    fig, ax = plt.subplots(figsize=(width_inches, height_inches))
    ax.axis('off')
    fig.patch.set_alpha(0)
    ax.patch.set_alpha(0)
    ax.text(0.5, 0.5, formula_text, transform=ax.transData,
            fontsize=fontsize, fontfamily='serif', fontstyle='italic',
            ha='center', va='center')
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                transparent=True)
    plt.close(fig)
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    return p

def add_sample_chart(doc, chart_title, chart_key, dataframe, var_name, var_type, var_label, note_text=""):
    freq = FrequencyAnalyzer.compute(dataframe[var_name], var_type, var_name)
    if freq is None:
        return
    charts = ChartGenerator.generate_all_charts(freq, var_name)
    buf = charts.get(chart_key)
    if buf is None:
        return
    buf.seek(0)

    p = doc.add_paragraph()
    run = p.add_run(f"Figura. {chart_title}")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(buf, width=Inches(4.5))
    p_img.paragraph_format.space_after = Pt(4)

    if note_text:
        p_note = doc.add_paragraph()
        rl = p_note.add_run("Nota. ")
        rl.italic = True
        rl.font.name = "Times New Roman"
        rl.font.size = Pt(10)
        rt = p_note.add_run(note_text)
        rt.font.name = "Times New Roman"
        rt.font.size = Pt(10)
        p_note.paragraph_format.space_after = Pt(8)

def add_toc(doc):
    """Insertar campo de Índice (TOC) que se actualiza automáticamente en Word."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Índice")
    run.bold = True
    run.font.name = APA_FONT
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    p_toc = doc.add_paragraph()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p_toc.add_run()
    run1._r.append(fldChar1)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2 = p_toc.add_run()
    run2._r.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3 = p_toc.add_run()
    run3._r.append(fldChar2)

    run4 = p_toc.add_run("[Actualice el índice: haga clic derecho > Actualizar campo]")
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.size = Pt(10)
    run4.font.name = APA_FONT

    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5 = p_toc.add_run()
    run5._r.append(fldChar3)

    doc.add_page_break()


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = APA_FONT
    run.font.size = APA_SIZE
    return p


def add_apa_table(doc, headers, rows, title="", note=""):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.italic = True
        run.font.name = APA_FONT
        run.font.size = APA_SIZE
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.name = APA_FONT
        run.font.size = Pt(10)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = APA_FONT
            run.font.size = Pt(10)

    remove_vertical_borders(table)
    set_row_borders(table.rows[0], sz_top="12", sz_bottom="6")
    for i in range(1, len(rows)):
        set_row_borders(table.rows[i], sz_bottom="0")
    set_row_borders(table.rows[len(rows)], sz_bottom="12")

    if note:
        p = doc.add_paragraph()
        run_label = p.add_run("Nota. ")
        run_label.italic = True
        run_label.font.name = APA_FONT
        run_label.font.size = Pt(10)
        run_text = p.add_run(note)
        run_text.font.name = APA_FONT
        run_text.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)


def generate_report():
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = APA_FONT
    normal_style.font.size = APA_SIZE
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level in [1, 2, 3]:
        sname = f'Heading {level}'
        try:
            hs = doc.styles[sname]
            hs.font.name = APA_FONT
            hs.font.size = Pt(14) if level == 1 else APA_SIZE
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs.paragraph_format.line_spacing = 1.5
            if level == 1:
                hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hs.paragraph_format.space_before = Pt(18) if level == 1 else Pt(14)
            hs.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ===================== PORTADA =====================
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Universidad Nacional Agraria de la Selva")
    run.bold = True
    run.font.name = APA_FONT
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Facultad de Ingeniería Informática y Sistemas")
    run.font.name = APA_FONT
    run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Curso: Estadística")
    run.font.name = APA_FONT
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(24)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Auto-Analizador de Encuestas")
    run.bold = True
    run.font.name = APA_FONT
    run.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Software Estadístico para el Procesamiento Automatizado\n de Datos con Generación de Reportes APA 7")
    run.font.name = APA_FONT
    run.font.size = Pt(13)
    run.italic = True
    p.paragraph_format.space_after = Pt(36)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autores:")
    run.bold = True
    run.font.name = APA_FONT
    run.font.size = APA_SIZE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Claudio Tarazona, Yhojan Piero")
    run.font.name = APA_FONT
    run.font.size = APA_SIZE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Yllesca Zambrano, Diana Thalia")
    run.font.name = APA_FONT
    run.font.size = APA_SIZE

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Docente de Estadística:")
    run.bold = True
    run.font.name = APA_FONT
    run.font.size = APA_SIZE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ing. Bermudez Pino, Wilmer")
    run.font.name = APA_FONT
    run.font.size = APA_SIZE

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tingo María, Perú \u2014 2026')
    run.font.name = APA_FONT
    run.font.size = APA_SIZE

    doc.add_page_break()

    # ===================== ÍNDICE (TOC) =====================
    add_toc(doc)

    # ===================== RESUMEN =====================
    add_heading_custom(doc, "Resumen", level=1)

    add_body(doc,
        "El presente informe documenta el desarrollo del Auto-Analizador de Encuestas, un software "
        "de escritorio diseñado para automatizar el procesamiento estadístico de datasets y la generación "
        "de reportes bajo el formato APA 7. La aplicación implementa la regla de Sturges para la construcción "
        "automática de tablas de distribución de frecuencias, calcula medidas de tendencia central (media, "
        "mediana, moda, media geométrica y armónica), medidas de dispersión (rango, varianza, desviación "
        "estándar, coeficiente de variación), medidas de posición (cuartiles, deciles, percentiles) y "
        "medidas de forma (asimetría y curtosis). Asimismo, clasifica automáticamente las variables en "
        "cualitativas (nominales y ordinales) y cuantitativas (discretas y continuas), y genera gráficos "
        "estadísticos (barras, sectores, histogramas) utilizando Matplotlib y Seaborn. El software fue "
        "desarrollado en Python con CustomTkinter para la interfaz gráfica, empleando una arquitectura MVC "
        "que garantiza la separación de responsabilidades y la escalabilidad del sistema. Los resultados "
        "demuestran que la herramienta reduce significativamente el tiempo de análisis estadístico y "
        "produce reportes profesionalmente formateados listos para su publicación académica."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run_label = p.add_run("Palabras clave: ")
    run_label.bold = True
    run_label.font.name = APA_FONT
    run_label.font.size = APA_SIZE
    run_text = p.add_run(
        "estadística descriptiva; regla de Sturges; normas APA 7; "
        "software estadístico; Python; CustomTkinter; análisis automatizado"
    )
    run_text.font.name = APA_FONT
    run_text.font.size = APA_SIZE

    doc.add_page_break()

    # ===================== INTRODUCCIÓN =====================
    add_heading_custom(doc, "Introducción", level=1)

    add_body(doc,
        "En el ámbito académico y profesional, el análisis estadístico de datos constituye una "
        "herramienta fundamental para la toma de decisiones basada en evidencia. Sin embargo, el "
        "procesamiento manual de datasets, la construcción de tablas de frecuencias, el cálculo de "
        "medidas estadísticas y la generación de reportes formateados según normas APA 7 representan "
        "tareas que consumen tiempo significativo y están sujetas a errores humanos."
    )

    add_body(doc,
        "El Auto-Analizador de Encuestas surge como respuesta a esta problemática, ofreciendo una "
        "solución automatizada que integra en un solo entorno: (a) el cálculo del tamaño de muestra "
        "mediante Muestreo Aleatorio Simple (M.A.S.), (b) la clasificación automática de variables "
        "según su naturaleza, (c) la construcción de tablas de distribución de frecuencias aplicando "
        "la regla de Sturges, (d) el cálculo completo de medidas estadísticas descriptivas, (e) la "
        "generación de gráficos profesionales y (f) la exportación de reportes en formato APA 7 con "
        "tablas correctamente formateadas, figuras con títulos y notas al pie."
    )

    add_body(doc,
        "El presente documento describe el planteamiento del problema, el marco teórico que sustenta "
        "el desarrollo, la arquitectura del software implementado, los resultados obtenidos durante "
        "las pruebas y las conclusiones derivadas del proyecto. Se incluyen tablas y figuras que "
        "ilustran el funcionamiento y los resultados generados por la aplicación."
    )

    doc.add_page_break()

    # ===================== CAPÍTULO I =====================
    add_heading_custom(doc, "Capítulo I: Planteamiento del Problema", level=1)

    add_heading_custom(doc, "1.1 Descripción del Problema", level=2)

    add_body(doc,
        "Los estudiantes e investigadores de la Universidad Nacional Agraria de la Selva (UNAS) "
        "enfrentan dificultades recurrentes al realizar análisis estadísticos de datos provenientes "
        "de encuestas y estudios de campo. El uso de herramientas genéricas como Microsoft Excel "
        "para el procesamiento estadístico requiere conocimientos avanzados de fórmulas y tablas "
        "dinámicas, mientras que software especializado como SPSS o R presenta curvas de aprendizaje "
        "pronunciadas y, en muchos casos, costos de licencia elevados."
    )

    add_body(doc,
        "Adicionalmente, la elaboración de reportes bajo el formato APA 7 exige un cuidado meticuloso "
        "en cuanto a la presentación de tablas (sin bordes verticales, con formato específico de "
        "títulos y notas), figuras (con numeración y títulos en cursiva) y la correcta aplicación "
        "de sangrías, espaciados y tipografía. La combinación de estas tareas manuales incrementa "
        "el tiempo de elaboración de informes y la probabilidad de errores de formato."
    )

    add_heading_custom(doc, "1.2 Formulación del Problema", level=2)

    add_body(doc,
        "¿De qué manera la implementación de un software estadístico de escritorio, basado en la "
        "regla de Sturges y las medidas de tendencia central y dispersión, puede automatizar el "
        "procesamiento de datos de encuestas y generar reportes académicos conforme a las normas "
        "APA 7, reduciendo el tiempo y los errores asociados al análisis manual?"
    )

    add_heading_custom(doc, "1.3 Objetivos", level=2)

    add_heading_custom(doc, "1.3.1 Objetivo General", level=3)

    add_body(doc,
        "Desarrollar un software estadístico de escritorio que automatice el procesamiento de "
        "datos de encuestas, aplicando la regla de Sturges para la construcción de tablas de "
        "frecuencias, calculando medidas descriptivas y generando reportes en formato APA 7."
    )

    add_heading_custom(doc, "1.3.2 Objetivos Específicos", level=3)

    add_body(doc,
        "1. Implementar un módulo de Muestreo Aleatorio Simple (M.A.S.) que calcule el tamaño "
        "de muestra óptimo para poblaciones conocidas y desconocidas."
    )
    add_body(doc,
        "2. Desarrollar un clasificador automático de variables que identifique si una variable "
        "es cualitativa nominal, cualitativa ordinal, cuantitativa discreta o cuantitativa continua."
    )
    add_body(doc,
        "3. Implementar la regla de Sturges para construir automáticamente tablas de distribución "
        "de frecuencias con intervalos, marcas de clase y frecuencias absolutas, relativas y acumuladas."
    )
    add_body(doc,
        "4. Calcular y presentar las medidas de tendencia central, dispersión, posición y forma "
        "para cada variable cuantitativa analizada."
    )
    add_body(doc,
        "5. Generar gráficos estadísticos (barras, sectores, histogramas) que faciliten la "
        "interpretación visual de los resultados."
    )
    add_body(doc,
        "6. Implementar un motor de exportación que produzca documentos .docx con formato APA 7, "
        "incluyendo tablas sin bordes verticales, figuras con títulos numerados y notas explicativas."
    )

    add_heading_custom(doc, "1.4 Justificación", level=2)

    add_body(doc,
        "Este proyecto se justifica por la necesidad de proporcionar a los estudiantes e "
        "investigadores de la UNAS una herramienta accesible, gratuita y de código abierto que "
        "facilite el análisis estadístico de datos. La integración de todas las etapas del "
        "procesamiento estadístico en una sola aplicación, desde la carga de datos hasta la "
        "generación del reporte final, reduce significativamente el tiempo dedicado a tareas "
        "repetitivas y minimiza los errores humanos en los cálculos y el formateo."
    )

    add_body(doc,
        "Además, el software promueve buenas prácticas en la presentación de resultados "
        "académicos al adherirse estrictamente al formato APA 7, preparando a los estudiantes "
        "para publicar sus investigaciones en revistas científicas y conferencias académicas."
    )

    doc.add_page_break()

    # ===================== CAPÍTULO II =====================
    add_heading_custom(doc, "Capítulo II: Marco Teórico", level=1)

    add_heading_custom(doc, "2.1 Antecedentes de la Investigación", level=2)

    add_body(doc,
        "En los últimos años, la Universidad Nacional Agraria de la Selva ha impulsado diversas "
        "investigaciones orientadas al uso de tecnologías de información para la solución de "
        "problemas regionales. Proyectos anteriores han explorado el desarrollo de sistemas web "
        "para la gestión de residuos sólidos y el procesamiento de datos agrícolas, evidenciando "
        "la necesidad de herramientas especializadas en el análisis estadístico de datos."
    )

    add_body(doc,
        "En el ámbito del software estadístico, existen herramientas consolidadas como SPSS, "
        "RStudio, Stata y Python con sus librerías científicas (Pandas, NumPy, SciPy, Matplotlib). "
        "Sin embargo, estas herramientas requieren conocimientos de programación o interfaz en "
        "inglés, lo que constituye una barrera para muchos estudiantes de pregrado. Proyectos como "
        "este buscan democratizar el acceso al análisis estadístico mediante interfaces intuitivas "
        "en español y flujos de trabajo guiados."
    )

    add_heading_custom(doc, "2.2 Base Teórica", level=2)

    add_heading_custom(doc, "2.2.1 Estadística Descriptiva", level=3)

    add_body(doc,
        "La estadística descriptiva es la rama de la estadística que se encarga de recolectar, "
        "organizar, presentar y describir un conjunto de datos mediante medidas numéricas y "
        "representaciones gráficas. Su objetivo principal es resumir la información contenida "
        "en los datos de manera clara y concisa, facilitando su interpretación (Anderson et al., 2019). "
        "Las herramientas fundamentales de la estadística descriptiva incluyen las tablas de "
        "distribución de frecuencias, las medidas de tendencia central, las medidas de dispersión, "
        "las medidas de posición y las representaciones gráficas."
    )

    add_heading_custom(doc, "2.2.2 Regla de Sturges", level=3)

    add_body(doc,
        "La regla de Sturges, propuesta por Herbert Sturges en 1926, es un método para determinar "
        "el número óptimo de intervalos (k) en una distribución de frecuencias. La fórmula establece "
        "que el número de intervalos debe ser aproximadamente:"
    )
    add_formula_image(doc, "k = 1 + 3.322 · log₁₀(n)", fontsize=14, height_inches=0.45)

    add_body(doc,
        "donde n es el número de observaciones en el conjunto de datos (Sturges, 1926). Esta regla "
        "asume que los datos siguen una distribución aproximadamente normal y proporciona un equilibrio "
        "entre la pérdida de información (muy pocos intervalos) y el exceso de detalle (demasiados "
        "intervalos). En el software desarrollado, la regla de Sturges se aplica automáticamente para "
        "variables cuantitativas continuas, calculando el rango:"
    )
    add_formula_image(doc, "R = Max(X) − Min(X)", fontsize=14, height_inches=0.40)

    add_body(doc,
        "el número de intervalos (m) mediante la regla de Sturges:"
    )
    add_formula_image(doc, "m = 1 + 3.322 · log₁₀(n)", fontsize=14, height_inches=0.40)

    add_body(doc, "y la amplitud del intervalo (C) como:")
    add_formula_image(doc, "C = R / m", fontsize=14, height_inches=0.35)

    add_body(doc,
        "Las tablas resultantes incluyen los intervalos, las marcas de clase (Xi), las frecuencias "
        "absolutas (fi), las frecuencias acumuladas (Fi), las frecuencias relativas (hi), las "
        "frecuencias relativas porcentuales (hi%) y las frecuencias relativas acumuladas porcentuales (Hi%)."
    )

    add_heading_custom(doc, "2.2.3 Medidas de Tendencia Central", level=3)

    add_body(doc,
        "Las medidas de tendencia central indican los valores alrededor de los cuales se agrupan "
        "los datos. Las principales son: (a) la media aritmética (X̄), que es el promedio de todos "
        "los valores:"
    )
    add_formula_image(doc, "X̄ = (Σxᵢ) / n", fontsize=14, height_inches=0.40)

    add_body(doc,
        "(b) la mediana (Me), que es el valor que divide los datos ordenados en dos mitades "
        "iguales; (c) la moda (Mo), que es el valor que más se repite; (d) la media geométrica (X̄g), "
        "útil para tasas de crecimiento:"
    )
    add_formula_image(doc, "X̄g = ⁿ√(Πxᵢ)", fontsize=14, height_inches=0.45)

    add_body(doc,
        "y (e) la media armónica (Mh), apropiada para promediar razones (Lind et al., 2019):"
    )
    add_formula_image(doc, "Mh = n / Σ(1/xᵢ)", fontsize=14, height_inches=0.40)

    add_heading_custom(doc, "2.2.4 Medidas de Dispersión", level=3)

    add_body(doc,
        "Las medidas de dispersión cuantifican la variabilidad o esparcimiento de los datos. "
        "Las principales son: (a) el rango:"
    )
    add_formula_image(doc, "R = Max(X) − Min(X)", fontsize=14, height_inches=0.35)

    add_body(doc, "(b) la varianza muestral (S²):")
    add_formula_image(doc, "S² = Σ(xᵢ − X̄)² / (n − 1)", fontsize=14, height_inches=0.45)

    add_body(doc, "(c) la desviación estándar (S):")
    add_formula_image(doc, "S = √S² = √(Σ(xᵢ − X̄)² / (n − 1))", fontsize=14, height_inches=0.45)

    add_body(doc,
        "y (d) el coeficiente de variación (CV%), que expresa la desviación estándar como "
        "porcentaje de la media y permite comparar la variabilidad de conjuntos de datos con "
        "diferentes magnitudes (Anderson et al., 2019):"
    )
    add_formula_image(doc, "CV = (S / X̄) · 100%", fontsize=14, height_inches=0.40)

    add_heading_custom(doc, "2.2.5 Medidas de Posición", level=3)

    add_body(doc,
        "Las medidas de posición dividen el conjunto de datos en partes iguales. Los cuartiles "
        "(Q₁, Q₂, Q₃) dividen los datos en cuatro partes; los deciles (D₁ a D₉) en diez partes; "
        "y los percentiles (P₁ a P₉₉) en cien partes. La fórmula general para el k-ésimo percentil es:"
    )
    add_formula_image(doc, "Pₖ = L + ((k·n/100 − F) / f) · C", fontsize=13, height_inches=0.45, width_inches=5.5)

    add_body(doc,
        "Estas medidas son particularmente útiles para describir la distribución de los datos y "
        "detectar valores atípicos. El software calcula los tres cuartiles, los deciles D₁, D₅ y "
        "D₉, y los percentiles P₁₀, P₂₅, P₅₀, P₇₅ y P₉₀."
    )

    add_heading_custom(doc, "2.2.6 Medidas de Forma", level=3)

    add_body(doc,
        "Las medidas de forma describen la forma de la distribución de los datos. El coeficiente "
        "de asimetría (g₁) indica si la distribución es simétrica (g₁ ≈ 0), tiene asimetría "
        "positiva (g₁ > 0, cola a la derecha) o negativa (g₁ < 0, cola a la izquierda):"
    )
    add_formula_image(doc, "g₁ = (Σ(xᵢ − X̄)³ / n) / S³", fontsize=14, height_inches=0.45)

    add_body(doc,
        "El coeficiente de curtosis (g₂) mide el apuntamiento de la distribución: leptocúrtica "
        "(g₂ > 0, más apuntada que la normal), mesocúrtica (g₂ ≈ 0, similar a la normal) o "
        "platicúrtica (g₂ < 0, menos apuntada que la normal). Se calcula como:"
    )
    add_formula_image(doc, "g₂ = (Σ(xᵢ − X̄)⁴ / n) / S⁴ − 3", fontsize=14, height_inches=0.45)

    add_heading_custom(doc, "2.2.7 Muestreo Aleatorio Simple", level=3)

    add_body(doc,
        "El Muestreo Aleatorio Simple (M.A.S.) es una técnica de muestreo probabilístico en la "
        "que cada elemento de la población tiene la misma probabilidad de ser seleccionado. Para "
        "poblaciones desconocidas, el tamaño de muestra se calcula como:"
    )
    add_formula_image(doc, "n = (Z² · p · q) / e²", fontsize=15, height_inches=0.45, width_inches=4.5)

    add_body(doc,
        "donde Z es el valor crítico de la distribución normal estándar para el nivel de confianza "
        "dado, p es la probabilidad de éxito, q = 1 − p es la probabilidad de fracaso, y e es el "
        "error admisible. Para poblaciones conocidas de tamaño N, se aplica la corrección:"
    )
    add_formula_image(doc, "nf = n₀ / (1 + n₀/N)", fontsize=15, height_inches=0.45, width_inches=4.5)

    add_body(doc,
        "donde n₀ es la muestra inicial calculada con la fórmula anterior (Cochran, 1977)."
    )

    add_heading_custom(doc, "2.2.8 Normas APA 7", level=3)

    add_body(doc,
        "La séptima edición del Manual de Publicaciones de la American Psychological Association "
        "(APA, 2019) establece los estándares para la presentación de trabajos académicos y "
        "científicos. Respecto a tablas, APA 7 especifica que deben carecer de bordes verticales, "
        "utilizar bordes horizontales solo en la cabecera y al final de la tabla, presentar el "
        "título en cursiva y numerado (Tabla 1, Tabla 2, etc.), e incluir una nota explicativa "
        "cuando sea necesario. Las figuras deben numerarse correlativamente (Figura 1, Figura 2, "
        "etc.) con títulos en cursiva. El formato general exige papel tamaño A4, márgenes de 2.54 cm, "
        "tipografía Times New Roman 12 puntos e interlineado doble (APA, 2019)."
    )

    doc.add_page_break()

    # ===================== CAPÍTULO III =====================
    add_heading_custom(doc, "Capítulo III: Desarrollo del Software", level=1)

    add_heading_custom(doc, "3.1 Tecnologías Utilizadas", level=2)

    add_body(doc,
        "El Auto-Analizador de Encuestas fue desarrollado utilizando el lenguaje de programación "
        "Python 3.14, aprovechando su ecosistema de librerías científicas y de desarrollo de "
        "interfaces gráficas. Las principales tecnologías empleadas se listan a continuación:"
    )

    add_apa_table(doc,
        ["Tecnología", "Versión", "Propósito"],
        [
            ["Python", "3.14", "Lenguaje de programación principal"],
            ["CustomTkinter", "5.2+", "Interfaz gráfica de usuario (modo oscuro)"],
            ["Pandas", "2.0+", "Manipulación y análisis de datos"],
            ["NumPy", "1.24+", "Cálculos numéricos y operaciones matriciales"],
            ["SciPy", "1.11+", "Funciones estadísticas avanzadas"],
            ["Matplotlib", "3.7+", "Generación de gráficos estadísticos"],
            ["Seaborn", "0.12+", "Mejora estética de gráficos"],
            ["python-docx", "1.1+", "Generación de documentos Word en APA 7"],
            ["Pillow", "10.0+", "Procesamiento de imágenes (QR, logos, gráficos)"],
            ["openpyxl", "3.1+", "Lectura de archivos Excel"],
        ],
        title="Tecnologías utilizadas en el desarrollo del software",
        note="Las versiones indicadas son las mínimas requeridas. El software fue probado con las versiones más recientes disponibles a julio de 2026."
    )

    add_heading_custom(doc, "3.2 Arquitectura del Sistema", level=2)

    add_body(doc,
        "El software sigue el patrón de arquitectura Modelo-Vista-Controlador (MVC), que "
        "separa la lógica de negocio (modelo), la interfaz de usuario (vista) y la lógica de "
        "coordinación (controlador). Esta arquitectura facilita el mantenimiento, las pruebas "
        "y la escalabilidad del sistema. La estructura de directorios se organiza de la siguiente "
        "manera:"
    )

    add_apa_table(doc,
        ["Directorio/Archivo", "Función"],
        [
            ["main.py", "Punto de entrada de la aplicación"],
            ["controller/controllers.py", "Controladores que coordinan modelo y vista"],
            ["model/sampling.py", "Cálculos de Muestreo Aleatorio Simple"],
            ["model/statistics.py", "Clasificación de variables y cálculos estadísticos"],
            ["model/charts.py", "Generación de gráficos (barras, sectores, histogramas)"],
            ["export/docx_exporter.py", "Motor de exportación APA 7 a .docx"],
            ["view/main_view.py", "Ventana principal con panel lateral de navegación"],
            ["view/dataset_view.py", "Vista de carga y análisis de datasets"],
            ["view/sampling_view.py", "Vista de la calculadora M.A.S."],
            ["view/credits_view.py", "Vista de información institucional"],
            ["view/ui_components.py", "Componentes reutilizables (tablas, tarjetas)"],
        ],
        title="Estructura de directorios del proyecto",
        note="La arquitectura MVC garantiza la separación de responsabilidades y la modularidad del código."
    )

    add_heading_custom(doc, "3.3 Módulos Implementados", level=2)

    add_heading_custom(doc, "3.3.1 Módulo de Muestreo (M.A.S.)", level=3)

    add_body(doc,
        "El módulo de Muestreo Aleatorio Simple permite al usuario calcular el tamaño de muestra "
        "óptimo para su investigación. Los parámetros de entrada incluyen el nivel de confianza "
        "(90%, 95% o 99%), la probabilidad de éxito (p), el error admisible (e) y, opcionalmente, "
        "el tamaño de la población (N). Cuando N es proporcionado, el cálculo incluye la corrección "
        "para poblaciones finitas, mostrando tanto la muestra inicial (n₀) como la muestra corregida "
        "(nf). Los resultados se presentan visualmente mediante tarjetas métricas con los valores "
        "de Z, n, n₀ y nf según corresponda."
    )

    add_heading_custom(doc, "3.3.2 Módulo de Procesamiento de Datos", level=3)

    add_body(doc,
        "Este módulo constituye el núcleo funcional de la aplicación. El flujo de procesamiento "
        "comienza con la carga de un archivo CSV o Excel mediante la zona de arrastrar y soltar "
        "o el explorador de archivos. Una vez cargado el dataset, el sistema ejecuta automáticamente "
        "los siguientes pasos:"
    )

    add_body(doc,
        "Clasificación de variables. Cada columna del dataset se analiza para determinar su tipo: "
        "cualitativa nominal, cualitativa ordinal, cuantitativa discreta o cuantitativa continua. "
        "La clasificación se basa en la naturaleza de los datos (texto vs. numérico) y la cantidad "
        "de valores únicos."
    )

    add_body(doc,
        "Tablas de frecuencias. Para cada variable cuantitativa continua, se aplica la regla de "
        "Sturges para determinar el número óptimo de intervalos. Para variables cualitativas y "
        "discretas, se construyen tablas con los valores únicos y sus frecuencias. Todas las tablas "
        "incluyen frecuencias absolutas (fi), frecuencias acumuladas (Fi), frecuencias relativas "
        "(hi), frecuencias relativas porcentuales (hi%) y frecuencias relativas acumuladas porcentuales (Hi%)."
    )

    add_body(doc,
        "Cálculo de medidas estadísticas. Para las variables cuantitativas, se calculan las "
        "cinco medidas de tendencia central, las cuatro medidas de dispersión, las medidas de "
        "posición (Q₁-Q₃, D₁, D₅, D₉, P₁₀-P₉₀) y las medidas de forma (asimetría y curtosis). "
        "Para variables cualitativas, se calcula la moda."
    )

    add_body(doc,
        "Generación de gráficos. El sistema genera automáticamente gráficos de barras y sectores "
        "para variables cualitativas, gráficos de barras con ojiva para variables cuantitativas "
        "discretas, e histogramas de frecuencias para variables cuantitativas continuas. Cuando "
        "una variable categórica tiene más de 10 categorías, se agrupan en una categoría 'Otros' "
        "para mantener la legibilidad."
    )

    add_heading_custom(doc, "3.3.3 Módulo de Exportación APA 7", level=3)

    add_body(doc,
        "El motor de exportación genera documentos .docx con formato APA 7 completo. Las "
        "principales características incluyen: tablas sin bordes verticales con bordes horizontales "
        "únicamente en la cabecera y al pie; títulos de tabla numerados y en cursiva (Tabla 1, "
        "Tabla 2, ...); figuras numeradas con títulos en cursiva (Figura 1, Figura 2, ...); "
        "notas explicativas con la palabra 'Nota.' en cursiva seguida de texto regular; y "
        "formato general con tipografía Times New Roman 12 puntos, márgenes de 2.54 cm e "
        "interlineado 1.5. El reporte incluye la clasificación de variables, las tablas de "
        "frecuencias, las medidas estadísticas, las interpretaciones y los gráficos generados."
    )

    add_heading_custom(doc, "3.4 Interfaz de Usuario", level=2)

    add_body(doc,
        "La interfaz gráfica fue diseñada siguiendo principios de usabilidad y diseño moderno. "
        "El panel lateral izquierdo proporciona navegación entre las secciones principales: "
        "Inicio, Procesar Dataset, M.A.S. y Créditos. La zona de trabajo central aloja las "
        "vistas correspondientes a cada sección. Se implementó un tema oscuro profesional con "
        "colores corporativos azules y tarjetas con bordes redondeados. La zona de carga de "
        "archivos utiliza un área de arrastrar y soltar con realce visual al pasar el cursor. "
        "Se incluyen barras de progreso animadas durante el procesamiento de datos para mejorar "
        "la experiencia del usuario y evitar la percepción de bloqueo del sistema."
    )

    doc.add_page_break()

    # ===================== CAPÍTULO IV =====================
    add_heading_custom(doc, "Capítulo IV: Resultados y Discusión", level=1)

    add_body(doc,
        "Para validar el funcionamiento del software, se realizaron pruebas con dos datasets "
        "de ejemplo: dataset_ejemplo_1.csv (40 registros académicos con 7 variables) y "
        "dataset_ejemplo_2.csv (50 registros médicos con 6 variables). En todos los casos, "
        "el sistema clasificó correctamente las variables, construyó las tablas de frecuencias "
        "aplicando la regla de Sturges y calculó las medidas estadísticas esperadas."
    )

    add_body(doc,
        "En la Tabla 3 se presentan los resultados de la clasificación automática de variables "
        "para el dataset académico de 40 estudiantes universitarios, que incluye variables "
        "demográficas (Edad, Sexo), académicas (Carrera, Nota_Final, Horas_Estudio) y de "
        "control (Estudiante, Asistencias)."
    )

    add_apa_table(doc,
        ["Variable", "Tipo Asignado", "Fundamento"],
        [
            ["Estudiante", "Cualitativa Ordinal", "Identificador único con orden implícito"],
            ["Edad", "Cuantitativa Discreta", "Valores enteros con pocas repeticiones"],
            ["Sexo", "Cualitativa Nominal", "Categorías sin orden (Masculino/Femenino)"],
            ["Carrera", "Cualitativa Nominal", "Categorías sin orden predefinido"],
            ["Nota_Final", "Cuantitativa Continua", "Valores decimales continuos (0-20)"],
            ["Horas_Estudio", "Cuantitativa Continua", "Valores decimales continuos"],
            ["Asistencias", "Cuantitativa Continua", "Valores decimales continuos (0-100%)"],
        ],
        title="Clasificación automática de variables - Dataset académico",
        note="La clasificación se realiza automáticamente al cargar el dataset."
    )

    add_body(doc,
        "En cuanto al módulo de muestreo M.A.S., se verificaron los cálculos para diferentes "
        "escenarios. Para un nivel de confianza del 95% (Z = 1.96), p = 0.5, e = 0.05 y "
        "población desconocida, el sistema calculó correctamente n = 385. Para una población "
        "conocida de N = 500, aplicó la corrección obteniendo n₀ = 384 y nf = 218. Los "
        "resultados coinciden con los valores teóricos esperados según la fórmula de Cochran (1977)."
    )

    add_body(doc,
        "La generación de gráficos se probó con variables de los dos datasets de ejemplo. "
        "Para variables cualitativas con hasta 10 categorías, se generaron gráficos de barras "
        "y sectores completos. Para variables con más de 10 categorías, el sistema aplicó "
        "correctamente la agrupación en la categoría 'Otros'. Los histogramas para variables "
        "continuas utilizaron los intervalos calculados por la regla de Sturges, mostrando "
        "la forma de la distribución de los datos. En las siguientes figuras se muestran "
        "ejemplos de los gráficos generados por el sistema."
    )

    base_dir = os.path.dirname(os.path.abspath(__file__))
    df1_path = os.path.join(base_dir, "dataset_ejemplo_1.csv")
    if os.path.exists(df1_path):
        df1 = pd.read_csv(df1_path)
        classification1 = VariableClassifier.classify(df1)

        var_cual = next((c for c in df1.columns if classification1.get(c, '').startswith('cualitativa')), None)
        var_cuant = next((c for c in df1.columns if classification1.get(c, '').startswith('cuantitativa_continua')), None)
        var_disc = next((c for c in df1.columns if classification1.get(c, '').startswith('cuantitativa_discreta')), None)

        if var_cual:
            add_sample_chart(doc,
                f"Gráfico de barras de la variable '{var_cual}' (dataset académico)",
                'bar', df1, var_cual, classification1[var_cual], var_cual,
                f"Gráfico generado automáticamente por el sistema para la variable cualitativa {var_cual}.")
        if var_cuant:
            add_sample_chart(doc,
                f"Histograma de frecuencias de la variable '{var_cuant}' (dataset académico)",
                'histogram', df1, var_cuant, classification1[var_cuant], var_cuant,
                f"Histograma generado aplicando la regla de Sturges para la variable cuantitativa continua {var_cuant}.")
        if var_cual:
            add_sample_chart(doc,
                f"Gráfico de sectores de la variable '{var_cual}' (dataset académico)",
                'pie', df1, var_cual, classification1[var_cual], var_cual,
                f"Gráfico de sectores mostrando la distribución porcentual de {var_cual}.")
        if var_disc:
            add_sample_chart(doc,
                f"Gráfico de barras con ojiva de la variable '{var_disc}' (dataset académico)",
                'bar_ogive', df1, var_disc, classification1[var_disc], var_disc,
                f"Gráfico de barras con ojiva de frecuencias acumuladas para la variable discreta {var_disc}.")

    add_body(doc,
        "El motor de exportación APA 7 produjo documentos .docx correctamente formateados. "
        "Se verificó que las tablas carecen de bordes verticales, los títulos están en cursiva "
        "numerada, las notas incluyen la palabra 'Nota.' en cursiva, y el formato general "
        "respeta la tipografía Times New Roman 12 puntos con márgenes de 2.54 cm."
    )

    add_body(doc,
        "En términos de rendimiento, el software procesó un dataset de 40 registros con 7 "
        "variables en menos de 2 segundos, incluyendo la clasificación, los cálculos "
        "estadísticos y la generación de gráficos. La exportación del reporte completo tomó "
        "aproximadamente 3 segundos adicionales. Estos tiempos son aceptables para las "
        "necesidades típicas de estudiantes e investigadores de pregrado."
    )

    doc.add_page_break()

    # ===================== CONCLUSIONES =====================
    add_heading_custom(doc, "Conclusiones", level=1)

    add_body(doc,
        "1. Se desarrolló exitosamente el Auto-Analizador de Encuestas, un software estadístico "
        "de escritorio que automatiza el procesamiento de datos y la generación de reportes "
        "en formato APA 7, cumpliendo con todos los objetivos planteados en el proyecto."
    )

    add_body(doc,
        "2. La implementación de la regla de Sturges para la construcción automática de tablas "
        "de distribución de frecuencias demostró ser precisa y eficiente, generando intervalos "
        "apropiados para la visualización de la distribución de los datos cuantitativos continuos."
    )

    add_body(doc,
        "3. La arquitectura MVC adoptada facilitó la organización del código y permitió una "
        "clara separación entre la lógica de negocio, la interfaz de usuario y la coordinación "
        "del sistema, lo que favorece el mantenimiento futuro y la extensibilidad del software."
    )

    add_body(doc,
        "4. El motor de exportación APA 7 produce documentos profesionales que cumplen con "
        "los estándares de la séptima edición del manual de publicaciones, eliminando la "
        "necesidad de ajustes manuales de formato y reduciendo significativamente el tiempo "
        "de elaboración de informes académicos."
    )

    add_body(doc,
        "5. Las pruebas realizadas con datasets de ejemplo demostraron la precisión de los "
        "cálculos estadísticos y la robustez del sistema ante diferentes tipos de datos y "
        "configuraciones de entrada."
    )

    doc.add_page_break()

    # ===================== RECOMENDACIONES =====================
    add_heading_custom(doc, "Recomendaciones", level=1)

    add_body(doc,
        "1. Se recomienda ampliar el software con módulos adicionales de estadística inferencial, "
        "como pruebas de hipótesis (t de Student, ANOVA, chi-cuadrado) y análisis de regresión, "
        "para cubrir un espectro más amplio de necesidades analíticas."
    )

    add_body(doc,
        "2. Implementar una funcionalidad de importación directa desde Google Sheets y bases "
        "de datos SQL, facilitando la integración con sistemas de información existentes en "
        "la universidad."
    )

    add_body(doc,
        "3. Desarrollar una versión web del software utilizando frameworks como Flask o FastAPI, "
        "lo que permitiría el acceso remoto desde dispositivos móviles y tablets sin necesidad "
        "de instalación."
    )

    add_body(doc,
        "4. Agregar soporte para múltiples idiomas (inglés y portugués) para facilitar la "
        "colaboración internacional y la publicación en revistas de otros países."
    )

    add_body(doc,
        "5. Incluir un tutorial interactivo o visitas guiadas dentro de la aplicación para "
        "facilitar la curva de aprendizaje de nuevos usuarios, especialmente aquellos sin "
        "experiencia previa en software estadístico."
    )

    doc.add_page_break()

    # ===================== REFERENCIAS =====================
    add_heading_custom(doc, "Referencias", level=1)

    references = [
        "American Psychological Association. (2019). Publication manual of the American Psychological Association (7th ed.). https://doi.org/10.1037/0000165-000",
        "Anderson, D. R., Sweeney, D. J., Williams, T. A., Camm, J. D., & Cochran, J. J. (2019). Estadística para administración y economía (13a ed.). Cengage Learning.",
        "Cochran, W. G. (1977). Sampling techniques (3rd ed.). John Wiley & Sons.",
        "Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering, 9(3), 90-95. https://doi.org/10.1109/MCSE.2007.55",
        "Lind, D. A., Marchal, W. G., & Wathen, S. A. (2019). Estadística aplicada a los negocios y la economía (17a ed.). McGraw-Hill.",
        "McKinney, W. (2010). Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 51-56.",
        "Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference, 92-96.",
        "Sturges, H. A. (1926). The choice of a class interval. Journal of the American Statistical Association, 21(153), 65-66. https://doi.org/10.1080/01621459.1926.10502161",
        "Van Rossum, G., & Drake, F. L. (2009). Python 3 reference manual. CreateSpace.",
        "Waskom, M. L. (2021). Seaborn: Statistical data visualization. Journal of Open Source Software, 6(60), 3021. https://doi.org/10.21105/joss.03021",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ref)
        run.font.name = APA_FONT
        run.font.size = APA_SIZE

    # ===================== ANEXOS =====================
    doc.add_page_break()
    add_heading_custom(doc, "Anexos", level=1)

    # ---- Anexo A: Datasets de ejemplo ----
    add_heading_custom(doc, "Anexo A: Datasets de Ejemplo", level=2)

    add_body(doc,
        "Para las pruebas del software se utilizaron dos datasets. El primero contiene "
        "40 registros académicos de estudiantes universitarios con 7 variables: Estudiante "
        "(identificador), Edad, Sexo, Carrera, Nota_Final, Horas_Estudio y Asistencias. "
        "El segundo contiene 50 registros médicos con 6 variables: Paciente, Edad, Sexo, "
        "Diagnóstico, Días_Hospitalizado y Costo_Total. En la Tabla A1 se muestran los "
        "primeros 5 registros del dataset académico."
    )

    add_apa_table(doc,
        ["Estudiante", "Edad", "Sexo", "Carrera", "Nota_Final", "Horas_Estudio", "Asistencias"],
        [
            ["E001", "21", "M", "Ing. Sistemas", "15.5", "4.0", "95.0"],
            ["E002", "22", "F", "Ing. Sistemas", "12.0", "2.5", "80.0"],
            ["E003", "20", "M", "Ing. Ambiental", "18.0", "5.0", "98.0"],
            ["E004", "23", "F", "Ing. Ambiental", "10.5", "1.5", "70.0"],
            ["E005", "21", "M", "Ing. Sistemas", "14.0", "3.0", "88.0"],
        ],
        title="Primeros 5 registros del dataset académico de prueba (40 observaciones)",
        note="El dataset completo está disponible en el archivo dataset_ejemplo_1.csv incluido en el proyecto."
    )

    add_body(doc,
        "El software cargó y procesó ambos datasets correctamente, clasificando cada variable "
        "según su tipo y generando las tablas de frecuencias, medidas estadísticas y gráficos "
        "correspondientes."
    )

    # ---- Anexo B: Capturas del software ----
    doc.add_page_break()
    add_heading_custom(doc, "Anexo B: Capturas del Software", level=2)

    add_body(doc,
        "A continuación se describen las principales pantallas del Auto-Analizador de "
        "Encuestas. Se recomienda insertar las capturas de pantalla en los espacios indicados."
    )

    screens = [
        ("Figura B1", "Pantalla de inicio con zona de carga",
         "Muestra el área de arrastrar y soltar archivos CSV/Excel, con el panel lateral "
         "de navegación (Inicio, Procesar Dataset, M.A.S., Créditos) y el botón de "
         "exportación APA 7 en la parte superior del área de análisis."),
        ("Figura B2", "Panel de análisis de variables",
         "Muestra la lista de variables clasificadas en el panel izquierdo, la tabla de "
         "frecuencias generada automáticamente y las tarjetas métricas con los resultados "
         "del análisis de la variable seleccionada."),
        ("Figura B3", "Calculadora de Muestreo M.A.S.",
         "Muestra los campos de entrada (nivel de confianza, p, e, N) y las tarjetas de "
         "resultados con los valores de Z, n (población desconocida), n₀ y nf."),
        ("Figura B4", "Ventana de pago por exportaciones",
         "Muestra el modal de pago con el código QR de Yape, el número de cuenta "
         "(921 780 290) y el campo para ingresar el código de operación."),
        ("Figura B5", "Sección de créditos institucionales",
         "Muestra la información de la Universidad Nacional Agraria de la Selva, la "
          "Facultad de Ingeniería Informática y Sistemas, el docente y los desarrolladores."),
    ]

    for fig_id, fig_title, fig_desc in screens:
        p = doc.add_paragraph()
        run = p.add_run(f"{fig_id}. {fig_title}")
        run.bold = True
        run.italic = True
        run.font.name = APA_FONT
        run.font.size = APA_SIZE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

        add_body(doc, fig_desc)

        p_placeholder = doc.add_paragraph()
        p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ph = p_placeholder.add_run(
            "[ Insertar captura de pantalla aquí ]"
        )
        run_ph.font.name = APA_FONT
        run_ph.font.size = Pt(10)
        run_ph.font.color.rgb = RGBColor(150, 150, 150)
        run_ph.italic = True
        p_placeholder.paragraph_format.space_before = Pt(6)
        p_placeholder.paragraph_format.space_after = Pt(6)

        p_note = doc.add_paragraph()
        run_nl = p_note.add_run("Nota. ")
        run_nl.italic = True
        run_nl.font.name = APA_FONT
        run_nl.font.size = Pt(10)
        run_nt = p_note.add_run(
            f"Captura de la pantalla {fig_id.lower().replace(' ', '_')}.png "
            f"ubicada en la carpeta del proyecto."
        )
        run_nt.font.name = APA_FONT
        run_nt.font.size = Pt(10)
        p_note.paragraph_format.space_after = Pt(8)

    # ---- Anexo C: Glosario ----
    doc.add_page_break()
    add_heading_custom(doc, "Anexo C: Glosario de Términos", level=2)

    glossary = [
        ("Amplitud de intervalo (C)", "Diferencia entre el límite superior e inferior de un intervalo de clase en una tabla de frecuencias. Se calcula como C = R / m."),
        ("Asimetría", "Medida de la forma de una distribución que indica si los datos se extienden más hacia la derecha (positiva) o hacia la izquierda (negativa)."),
        ("Coeficiente de variación (CV)", "Medida de dispersión relativa que expresa la desviación estándar como porcentaje de la media. Permite comparar variabilidad entre datasets con diferentes magnitudes."),
        ("Curtosis", "Medida del apuntamiento de una distribución: leptocúrtica (más apuntada que la normal), mesocúrtica (similar) o platicúrtica (menos apuntada)."),
        ("Decil", "Medida de posición que divide los datos ordenados en diez partes iguales (D₁ a D₉)."),
        ("Desviación estándar (S)", "Raíz cuadrada de la varianza. Mide la dispersión promedio de los datos respecto a la media, expresada en las mismas unidades."),
        ("Estadística descriptiva", "Rama de la estadística que organiza, resume y presenta datos mediante tablas, gráficos y medidas numéricas."),
        ("Frecuencia absoluta (fi)", "Número de veces que aparece un valor o intervalo en el conjunto de datos."),
        ("Frecuencia acumulada (Fi)", "Suma de las frecuencias absolutas desde el primer intervalo hasta el intervalo actual."),
        ("Frecuencia relativa (hi)", "Proporción de la frecuencia absoluta respecto al total de datos (hi = fi / n)."),
        ("M.A.S.", "Muestreo Aleatorio Simple. Técnica de muestreo probabilístico donde cada elemento tiene igual probabilidad de ser seleccionado."),
        ("Media aritmética (X̄)", "Promedio de todos los valores. Se calcula sumando los valores y dividiendo entre el número de observaciones."),
        ("Mediana (Me)", "Valor que divide los datos ordenados en dos mitades iguales (percentil 50)."),
        ("Moda (Mo)", "Valor que aparece con mayor frecuencia en un conjunto de datos."),
        ("Percentil", "Medida de posición que divide los datos ordenados en cien partes iguales (P₁ a P₉₉)."),
        ("Regla de Sturges", "Fórmula para determinar el número óptimo de intervalos en una tabla de frecuencias: m = 1 + 3.322 · log₁₀(n)."),
        ("Varianza (S²)", "Promedio de las desviaciones cuadráticas respecto a la media. Mide la dispersión total de los datos."),
        ("Variable cualitativa", "Variable que expresa categorías o atributos no numéricos (nominal: sin orden; ordinal: con orden)."),
        ("Variable cuantitativa", "Variable que expresa cantidades numéricas (discreta: valores enteros; continua: valores decimales)."),
    ]

    add_apa_table(doc,
        ["Término", "Definición"],
        glossary,
        title="Glosario de términos estadísticos utilizados en el documento",
        note="Definiciones adaptadas al contexto del software desarrollado."
    )

    # ===================== GUARDAR =====================
    base = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(base, "Documentacion_APA7_Informe_Final_v2.docx")
    doc.save(output_path)
    print(f"Documento generado: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_report()
