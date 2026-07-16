# -*- coding: utf-8 -*-
"""Genera el informe final del proyecto en formato APA 7."""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Documentacion_APA7_Informe_Final_v3.docx')

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))

# ---------- helpers ----------

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_number(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def set_normal_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.27)

def add_heading_apa(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Cm(0)
        return p
    elif level == 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
        return p
    elif level == 3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
        return p
    elif level == 4:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
        return p
    return doc.add_paragraph(text)

def add_paragraph(doc, text, bold=False, italic=False, indent=True, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if align:
        p.alignment = align
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def make_table(doc, headers, rows, title=None, note=None, col_widths=None):
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        set_cell_shading(cell, "2B2B2B")
        run.font.color.rgb = RGBColor(255, 255, 255)
    # body
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
    if note:
        p = doc.add_paragraph()
        run = p.add_run('Nota. ')
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run2 = p.add_run(note)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(2)
    doc.add_paragraph()
    return table


# ======= Construction =======

doc = Document()

# margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

set_normal_style(doc)

# ---- PORTADA ----
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Universidad Nacional Agraria de la Selva')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
p.paragraph_format.first_line_indent = Cm(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Facultad de Ingeniería Informática y Sistemas')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Curso: Estadística')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Auto-Analizador de Encuestas')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
p.paragraph_format.first_line_indent = Cm(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Software Estadístico para el Procesamiento Automatizado\n'
                'de Datos con Generación de Reportes APA 7')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Autores:')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Claudio Tarazona, Yhojan Piero\n'
                'Yllesca Zambrano, Diana Thalia')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Docente de Estadística:')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Ing. Bermudez Pino, Wilmer')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tingo María, Perú — 2026')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.first_line_indent = Cm(0)

doc.add_page_break()

# ---- ÍNDICE ----
add_heading_apa(doc, 'Índice', 1)
toc_items = [
    ('Resumen', 3),
    ('Introducción', 4),
    ('Capítulo I: Planteamiento del Problema', 5),
    ('    1.1 Descripción del Problema', 5),
    ('    1.2 Formulación del Problema', 6),
    ('    1.3 Objetivos', 6),
    ('    1.4 Justificación', 7),
    ('Capítulo II: Marco Teórico', 8),
    ('    2.1 Antecedentes de la Investigación', 8),
    ('    2.2 Base Teórica', 9),
    ('    2.3 Ingeniería de Requisitos', 13),
    ('Capítulo III: Desarrollo del Software', 15),
    ('    3.1 Tecnologías Utilizadas', 15),
    ('    3.2 Arquitectura del Sistema', 16),
    ('    3.3 Módulos Implementados', 17),
    ('    3.4 Interfaz de Usuario', 20),
    ('Capítulo IV: Resultados y Discusión', 21),
    ('Conclusiones', 24),
    ('Recomendaciones', 25),
    ('Referencias', 26),
    ('Anexos', 28),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{item} ............................. {page}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)

doc.add_page_break()

# ---- RESUMEN ----
add_heading_apa(doc, 'Resumen', 1)

add_paragraph(doc,
    'El presente informe documenta el desarrollo del Auto-Analizador de Encuestas, un software '
    'estadístico que automatiza el procesamiento de datasets y la generación de reportes bajo el '
    'formato APA 7. La aplicación implementa la regla de Sturges para la construcción automática '
    'de tablas de distribución de frecuencias, calcula medidas de tendencia central (media, '
    'mediana, moda, media geométrica y armónica), medidas de dispersión (rango, varianza, '
    'desviación estándar, coeficiente de variación), medidas de posición (cuartiles, deciles, '
    'percentiles) y medidas de forma (asimetría y curtosis). Asimismo, clasifica automáticamente '
    'las variables en cualitativas (nominales y ordinales) y cuantitativas (discretas y continuas), '
    'y genera gráficos estadísticos (barras, sectores, histogramas, polígonos de frecuencia) '
    'utilizando Matplotlib y Seaborn. El software fue desarrollado en Python con dos interfaces: '
    'una aplicación de escritorio con CustomTkinter bajo una arquitectura MVC, y una aplicación '
    'web con Flask que incorpora un sistema completo de calificación y muro de reseñas anónimas. '
    'Los resultados demuestran que la herramienta reduce significativamente el tiempo de análisis '
    'estadístico y produce reportes profesionalmente formateados listos para su publicación académica.')

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('Palabras clave: ')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p.add_run('estadística descriptiva; regla de Sturges; normas APA 7; software estadístico; '
                  'Python; CustomTkinter; Flask; análisis automatizado; ingeniería de requisitos')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

doc.add_page_break()

# ---- INTRODUCCIÓN ----
add_heading_apa(doc, 'Introducción', 1)

add_paragraph(doc,
    'En el ámbito académico y profesional, el análisis estadístico de datos constituye una '
    'herramienta fundamental para la toma de decisiones basada en evidencia. Sin embargo, el '
    'procesamiento manual de datasets, la construcción de tablas de frecuencias, el cálculo de '
    'medidas estadísticas y la generación de reportes formateados según normas APA 7 representan '
    'tareas que consumen tiempo significativo y están sujetas a errores humanos.')

add_paragraph(doc,
    'El Auto-Analizador de Encuestas surge como respuesta a esta problemática, ofreciendo una '
    'solución automatizada que integra en un solo entorno: (a) el cálculo del tamaño de muestra '
    'mediante Muestreo Aleatorio Simple (M.A.S.), (b) la clasificación automática de variables '
    'según su naturaleza, (c) la construcción de tablas de distribución de frecuencias aplicando '
    'la regla de Sturges, (d) el cálculo completo de medidas estadísticas descriptivas, (e) la '
    'generación de gráficos profesionales y (f) la exportación de reportes en formato APA 7 con '
    'tablas correctamente formateadas, figuras con títulos y notas al pie.')

add_paragraph(doc,
    'El desarrollo del proyecto se abordó desde la perspectiva de la ingeniería de requisitos, '
    'aplicando técnicas de elicitación como entrevistas, observación directa y análisis de '
    'documentos para identificar las necesidades de los estudiantes e investigadores de la '
    'Universidad Nacional Agraria de la Selva (UNAS). Los requisitos funcionales y no funcionales '
    'obtenidos fueron especificados, modelados mediante casos de uso y validados a través de '
    'prototipado y pruebas de usabilidad.')

add_paragraph(doc,
    'El presente documento describe el planteamiento del problema, el marco teórico que sustenta '
    'el desarrollo —incluyendo los fundamentos de la ingeniería de requisitos aplicados—, la '
    'arquitectura del software implementado, los resultados obtenidos durante las pruebas y las '
    'conclusiones derivadas del proyecto. Se incluyen tablas y figuras que ilustran el '
    'funcionamiento y los resultados generados por la aplicación.')

doc.add_page_break()

# ================== CAPÍTULO I ==================
add_heading_apa(doc, 'Capítulo I: Planteamiento del Problema', 1)

add_heading_apa(doc, '1.1 Descripción del Problema', 2)

add_paragraph(doc,
    'Los estudiantes e investigadores de la Universidad Nacional Agraria de la Selva (UNAS) '
    'enfrentan dificultades recurrentes al realizar análisis estadísticos de datos provenientes '
    'de encuestas y estudios de campo. El uso de herramientas genéricas como Microsoft Excel '
    'para el procesamiento estadístico requiere conocimientos avanzados de fórmulas y tablas '
    'dinámicas, mientras que software especializado como SPSS o R presenta curvas de aprendizaje '
    'pronunciadas y, en muchos casos, costos de licencia elevados.')

add_paragraph(doc,
    'Adicionalmente, la elaboración de reportes bajo el formato APA 7 exige un cuidado meticuloso '
    'en cuanto a la presentación de tablas (sin bordes verticales, con formato específico de '
    'títulos y notas), figuras (con numeración y títulos en cursiva) y la correcta aplicación de '
    'sangrías, espaciados y tipografía. La combinación de estas tareas manuales incrementa el '
    'tiempo de elaboración de informes y la probabilidad de errores de formato.')

add_paragraph(doc,
    'Desde la perspectiva de la ingeniería de requisitos, se identificó que las necesidades '
    'expresadas por los usuarios —como "que sea rápido", "que sea fácil de usar" o "que no se '
    'pierda información"— constituían requisitos de alto nivel que requerían un proceso '
    'sistemático de elicitación y especificación para transformarlos en requisitos funcionales y '
    'no funcionales verificables (IEEE, 2014; ISO/IEC 25010:2023).')

add_heading_apa(doc, '1.2 Formulación del Problema', 2)

add_paragraph(doc,
    '¿De qué manera la implementación de un software estadístico, basado en la regla de Sturges '
    'y las medidas de tendencia central y dispersión, puede automatizar el procesamiento de '
    'datos de encuestas y generar reportes académicos conforme a las normas APA 7, reduciendo '
    'el tiempo y los errores asociados al análisis manual?')

add_heading_apa(doc, '1.3 Objetivos', 2)
add_heading_apa(doc, '1.3.1 Objetivo General', 3)

add_paragraph(doc,
    'Desarrollar un software estadístico que automatice el procesamiento de datos de encuestas, '
    'aplicando la regla de Sturges para la construcción de tablas de frecuencias, calculando '
    'medidas descriptivas y generando reportes en formato APA 7.')

add_heading_apa(doc, '1.3.2 Objetivos Específicos', 3)

add_paragraph(doc,
    '1. Implementar un módulo de Muestreo Aleatorio Simple (M.A.S.) que calcule el tamaño de '
    'muestra óptimo para poblaciones conocidas y desconocidas.')

add_paragraph(doc,
    '2. Desarrollar un clasificador automático de variables que identifique si una variable es '
    'cualitativa nominal, cualitativa ordinal, cuantitativa discreta o cuantitativa continua.')

add_paragraph(doc,
    '3. Implementar la regla de Sturges para construir automáticamente tablas de distribución '
    'de frecuencias con intervalos, marcas de clase y frecuencias absolutas, relativas y '
    'acumuladas.')

add_paragraph(doc,
    '4. Calcular y presentar las medidas de tendencia central, dispersión, posición y forma '
    'para cada variable cuantitativa analizada.')

add_paragraph(doc,
    '5. Generar gráficos estadísticos (barras, sectores, histogramas) que faciliten la '
    'interpretación visual de los resultados.')

add_paragraph(doc,
    '6. Implementar un motor de exportación que produzca documentos .docx con formato APA 7, '
    'incluyendo tablas sin bordes verticales, figuras con títulos numerados y notas explicativas.')

add_paragraph(doc,
    '7. Aplicar la metodología de ingeniería de requisitos para la elicitación, especificación '
    'y validación de requisitos funcionales y no funcionales del sistema, garantizando la '
    'trazabilidad entre artefactos del proyecto.')

add_heading_apa(doc, '1.4 Justificación', 2)

add_paragraph(doc,
    'Este proyecto se justifica por la necesidad de proporcionar a los estudiantes e '
    'investigadores de la UNAS una herramienta accesible, gratuita y de código abierto que '
    'facilite el análisis estadístico de datos. La integración de todas las etapas del '
    'procesamiento estadístico en una sola aplicación, desde la carga de datos hasta la '
    'generación del reporte final, reduce significativamente el tiempo dedicado a tareas '
    'repetitivas y minimiza los errores humanos en los cálculos y el formateo.')

add_paragraph(doc,
    'Además, el software promueve buenas prácticas en la presentación de resultados académicos '
    'al adherirse estrictamente al formato APA 7, preparando a los estudiantes para publicar '
    'sus investigaciones en revistas científicas y conferencias académicas.')

add_paragraph(doc,
    'Desde la óptica de la ingeniería de requisitos, el proyecto constituye un caso aplicativo '
    'completo donde se implementan las buenas prácticas de elicitación, especificación, modelado '
    'y validación establecidas en el estándar IEEE 830 y el modelo de calidad ISO/IEC 25010:2023, '
    'lo que permite obtener un producto de software con requisitos bien definidos, verificables '
    'y alineados con las necesidades reales de los usuarios.')

doc.add_page_break()

# ================== CAPÍTULO II ==================
add_heading_apa(doc, 'Capítulo II: Marco Teórico', 1)

add_heading_apa(doc, '2.1 Antecedentes de la Investigación', 2)

add_paragraph(doc,
    'En los últimos años, la Universidad Nacional Agraria de la Selva ha impulsado diversas '
    'investigaciones orientadas al uso de tecnologías de información para la solución de '
    'problemas regionales. Proyectos anteriores han explorado el desarrollo de sistemas web '
    'para la gestión de residuos sólidos y el procesamiento de datos agrícolas, evidenciando '
    'la necesidad de herramientas especializadas en el análisis estadístico de datos.')

add_paragraph(doc,
    'En el ámbito del software estadístico, existen herramientas consolidadas como SPSS, '
    'RStudio, Stata y Python con sus librerías científicas (Pandas, NumPy, SciPy, Matplotlib). '
    'Sin embargo, estas herramientas requieren conocimientos de programación o interfaz en '
    'inglés, lo que constituye una barrera para muchos estudiantes de pregrado. Proyectos como '
    'este buscan democratizar el acceso al análisis estadístico mediante interfaces intuitivas '
    'en español y flujos de trabajo guiados (McKinney, 2010; Hunter, 2007).')

add_heading_apa(doc, '2.2 Base Teórica', 2)

add_heading_apa(doc, '2.2.1 Estadística Descriptiva', 3)

add_paragraph(doc,
    'La estadística descriptiva es la rama de la estadística que se encarga de recolectar, '
    'organizar, presentar y describir un conjunto de datos mediante medidas numéricas y '
    'representaciones gráficas. Su objetivo principal es resumir la información contenida en '
    'los datos de manera clara y concisa, facilitando su interpretación (Anderson et al., 2019). '
    'Las herramientas fundamentales de la estadística descriptiva incluyen las tablas de '
    'distribución de frecuencias, las medidas de tendencia central, las medidas de dispersión, '
    'las medidas de posición y las representaciones gráficas. El software implementa todas '
    'estas herramientas en un flujo automatizado.')

add_heading_apa(doc, '2.2.2 Clasificación de Variables', 3)

add_paragraph(doc,
    'Las variables estadísticas se clasifican según su naturaleza en dos grandes grupos. '
    'Las variables cualitativas expresan categorías o atributos no numéricos y pueden ser '
    'nominales (sin orden inherente, como el sexo o la carrera profesional) u ordinales '
    '(con un orden implícito, como el nivel de satisfacción). Las variables cuantitativas '
    'expresan cantidades numéricas y pueden ser discretas (valores enteros, como el número '
    'de hijos) o continuas (valores decimales, como la altura o el peso).')

add_paragraph(doc,
    'El sistema implementa un clasificador automático basado en los tipos de datos del '
    'DataFrame de Pandas: las columnas de tipo float64 se clasifican como cuantitativas '
    'continuas, las de tipo int64 como cuantitativas discretas, y las de tipo object o '
    'string como cualitativas nominales. El usuario puede reclasificar manualmente cualquier '
    'variable a través de un menú desplegable.')

add_heading_apa(doc, '2.2.3 Regla de Sturges', 3)

add_paragraph(doc,
    'La regla de Sturges, propuesta por Herbert Sturges en 1926, es un método para determinar '
    'el número óptimo de intervalos (k) en una distribución de frecuencias. La fórmula '
    'establece que el número de intervalos debe ser aproximadamente:')

add_paragraph(doc, 'm = 1 + 3.322 · log₁₀(n)', indent=False, italic=True)

add_paragraph(doc,
    'donde n es el número de observaciones en el conjunto de datos (Sturges, 1926). Esta regla '
    'asume que los datos siguen una distribución aproximadamente normal y proporciona un '
    'equilibrio entre la pérdida de información (muy pocos intervalos) y el exceso de detalle '
    '(demasiados intervalos). En el software desarrollado, la regla de Sturges se aplica '
    'automáticamente para variables cuantitativas continuas y para variables discretas con más '
    'de 15 valores únicos. El proceso calcula el rango R = max − min, el número de intervalos '
    'm mediante la regla de Sturges, y la amplitud del intervalo C = ⌈R / m⌉. Las tablas '
    'resultantes incluyen los intervalos, las marcas de clase (Xi), las frecuencias absolutas '
    '(fi), acumuladas (Fi), relativas (hi), relativas porcentuales (hi%) y relativas '
    'acumuladas porcentuales (Hi%).')

add_heading_apa(doc, '2.2.4 Medidas de Tendencia Central', 3)

add_paragraph(doc,
    'Las medidas de tendencia central indican los valores alrededor de los cuales se agrupan '
    'los datos. El software calcula cinco medidas: (a) la media aritmética (X̄), que es el '
    'promedio de todos los valores; (b) la mediana (Me), que es el valor que divide los datos '
    'ordenados en dos mitades iguales; (c) la moda (Mo), que es el valor que más se repite; '
    '(d) la media geométrica (X̄g), útil para tasas de crecimiento; y (e) la media armónica '
    '(Mh), apropiada para promediar razones (Lind et al., 2019).')

add_heading_apa(doc, '2.2.5 Medidas de Dispersión', 3)

add_paragraph(doc,
    'Las medidas de dispersión cuantifican la variabilidad o esparcimiento de los datos. '
    'El software calcula: (a) el rango (R = max − min); (b) la varianza muestral (S²); '
    '(c) la desviación estándar (S); y (d) el coeficiente de variación (CV%), que expresa '
    'la desviación estándar como porcentaje de la media y permite comparar la variabilidad '
    'de conjuntos de datos con diferentes magnitudes (Anderson et al., 2019).')

add_heading_apa(doc, '2.2.6 Medidas de Posición', 3)

add_paragraph(doc,
    'Las medidas de posición dividen el conjunto de datos en partes iguales. Los cuartiles '
    '(Q₁, Q₂, Q₃) dividen los datos en cuatro partes; los deciles (D₁ a D₉) en diez partes; '
    'y los percentiles (P₁ a P₉₉) en cien partes. El software calcula los tres cuartiles, '
    'los deciles D₁, D₅ y D₉, y los percentiles P₁₀, P₂₅, P₅₀, P₇₅ y P₉₀, utilizando la '
    'fórmula general para el k-ésimo percentil.')

add_heading_apa(doc, '2.2.7 Medidas de Forma', 3)

add_paragraph(doc,
    'Las medidas de forma describen la forma de la distribución de los datos. El coeficiente '
    'de asimetría (g₁) indica si la distribución es simétrica (g₁ ≈ 0), tiene asimetría '
    'positiva (g₁ > 0, cola a la derecha) o negativa (g₁ < 0, cola a la izquierda). El '
    'coeficiente de curtosis (g₂) mide el apuntamiento de la distribución: leptocúrtica '
    '(g₂ > 0, más apuntada que la normal), mesocúrtica (g₂ ≈ 0, similar a la normal) o '
    'platicúrtica (g₂ < 0, menos apuntada que la normal).')

add_heading_apa(doc, '2.2.8 Muestreo Aleatorio Simple', 3)

add_paragraph(doc,
    'El Muestreo Aleatorio Simple (M.A.S.) es una técnica de muestreo probabilístico en la '
    'que cada elemento de la población tiene la misma probabilidad de ser seleccionado. Para '
    'poblaciones desconocidas, el tamaño de muestra se calcula como:')

add_paragraph(doc, 'n = (Z² · p · q) / e²', indent=False, italic=True)

add_paragraph(doc,
    'donde Z es el valor crítico de la distribución normal estándar para el nivel de '
    'confianza dado (1.645 para 90%, 1.960 para 95%, 2.576 para 99%), p es la probabilidad '
    'de éxito, q = 1 − p, y e es el error admisible. Para poblaciones conocidas de tamaño '
    'N, se aplica la corrección para poblaciones finitas: nf = n₀ / (1 + n₀/N), donde n₀ es '
    'la muestra inicial calculada con la fórmula anterior (Cochran, 1977).')

add_heading_apa(doc, '2.2.9 Normas APA 7 para Tablas y Figuras', 3)

add_paragraph(doc,
    'La séptima edición del Manual de Publicaciones de la American Psychological Association '
    '(APA, 2019) establece los estándares para la presentación de trabajos académicos y '
    'científicos. Respecto a tablas, APA 7 especifica que deben carecer de bordes verticales, '
    'utilizar bordes horizontales solo en la cabecera y al final de la tabla, presentar el '
    'título en cursiva y numerado (Tabla 1, Tabla 2, etc.), e incluir una nota explicativa '
    'cuando sea necesario. Las figuras deben numerarse correlativamente (Figura 1, Figura 2 '
    'etc.) con títulos en cursiva. El formato general exige papel tamaño A4, márgenes de '
    '2.54 cm, tipografía Times New Roman 12 puntos e interlineado doble (APA, 2019).')

add_heading_apa(doc, '2.2.10 Generación de Gráficos Estadísticos', 3)

add_paragraph(doc,
    'El software genera gráficos estadísticos utilizando las librerías Matplotlib (Hunter, '
    '2007) y Seaborn (Waskom, 2021). Para variables cualitativas se generan gráficos de '
    'barras y de sectores (pastel). Cuando una variable categórica tiene más de 10 categorías '
    '(o más de 8 para sectores), las categorías menos frecuentes se agrupan en una categoría '
    '"Otros" para mantener la legibilidad. Para variables cuantitativas discretas se genera '
    'un gráfico de barras con ojiva de frecuencias acumuladas superpuesta (eje dual). Para '
    'variables cuantitativas continuas se genera un histograma de frecuencias y un gráfico '
    'de polígono de frecuencia con ojiva.')

# ---- 2.3 Ingeniería de Requisitos ----
add_heading_apa(doc, '2.3 Ingeniería de Requisitos', 2)

add_paragraph(doc,
    'La ingeniería de requisitos es el proceso sistemático de desarrollar requisitos de '
    'software mediante un proceso iterativo y coordinado de análisis, documentación, '
    'revisión y evaluación de cambios. Su objetivo principal es garantizar que el producto '
    'final satisfaga las necesidades reales de los stakeholders (IEEE, 2014). El presente '
    'proyecto aplicó las buenas prácticas de la ingeniería de requisitos siguiendo las '
    'directrices del estándar IEEE 830 y el modelo de calidad ISO/IEC 25010:2023.')

add_heading_apa(doc, '2.3.1 Conceptos Fundamentales', 3)

add_paragraph(doc,
    'Un requisito es una condición o capacidad que debe cumplir un sistema para resolver '
    'un problema del cliente o alcanzar un objetivo. Los requisitos se clasifican en '
    'requisitos funcionales (RF), que describen las funciones específicas que el sistema '
    'debe realizar, y requisitos no funcionales (RNF), que describen restricciones y '
    'cualidades del sistema como rendimiento, seguridad, usabilidad y mantenibilidad.')

add_heading_apa(doc, '2.3.2 Criterios de Calidad de los Requisitos', 3)

add_paragraph(doc,
    'Según el estándar IEEE 830, un requisito de calidad debe cumplir con los siguientes '
    'criterios: completo (incluye toda la información necesaria), correcto (refleja '
    'fielmente la necesidad del negocio), unívoco (tiene una sola interpretación posible), '
    'verificable (se puede comprobar mediante prueba o inspección), consistente (no '
    'contradice a otros requisitos), trazable (se puede rastrear hasta su origen y '
    'artefactos), y priorizado (tiene prioridad asignada).')

add_heading_apa(doc, '2.3.3 Técnicas de Elicitación', 3)

add_paragraph(doc,
    'La elicitación de requisitos es el proceso de descubrir y obtener los requisitos '
    'del sistema a partir de los stakeholders. En este proyecto se aplicaron las siguientes '
    'técnicas: (a) entrevistas con estudiantes e investigadores de la UNAS para comprender '
    'sus necesidades de análisis estadístico; (b) observación directa del proceso manual '
    'de elaboración de informes estadísticos; (c) análisis de documentos como trabajos de '
    'investigación previos, formatos de informes y guías de estilo APA 7; y (d) encuestas '
    'para priorizar funcionalidades deseadas.')

add_heading_apa(doc, '2.3.4 Modelado de Requisitos', 3)

add_paragraph(doc,
    'El modelado de requisitos permite representar visualmente el comportamiento del sistema '
    'y las interacciones entre sus componentes. Las técnicas de modelado aplicadas incluyen '
    'diagramas de casos de uso (que muestran las interacciones entre los actores y el sistema), '
    'diagramas de contexto (que delimitan el alcance del sistema), historias de usuario '
    '(descripciones breves de funcionalidades desde la perspectiva del usuario) y prototipos '
    '(representaciones visuales de la interfaz para validación temprana).')

add_heading_apa(doc, '2.3.5 Validación de Requisitos', 3)

add_paragraph(doc,
    'La validación de requisitos asegura que los requisitos especificados reflejan '
    'correctamente las necesidades de los stakeholders. Las técnicas de validación aplicadas '
    'incluyen walkthroughs (revisión en grupo de los requisitos con los stakeholders), '
    'prototipado (validación de la interfaz mediante prototipos de baja fidelidad), '
    'pruebas de usabilidad (evaluación con usuarios reales), inspección formal (revisión '
    'sistemática con checklist de calidad) y pruebas de rendimiento (verificación de '
    'tiempos de respuesta bajo carga).')

doc.add_page_break()

# ================== CAPÍTULO III ==================
add_heading_apa(doc, 'Capítulo III: Desarrollo del Software', 1)

add_heading_apa(doc, '3.1 Tecnologías Utilizadas', 2)

add_paragraph(doc,
    'El Auto-Analizador de Encuestas fue desarrollado utilizando el lenguaje de programación '
    'Python 3.14, aprovechando su ecosistema de librerías científicas y de desarrollo de '
    'interfaces gráficas. Las principales tecnologías empleadas se listan en la Tabla 1.')

make_table(doc,
    ['Tecnología', 'Versión', 'Propósito'],
    [
        ['Python', '3.14', 'Lenguaje de programación principal'],
        ['CustomTkinter', '5.2+', 'Interfaz gráfica de escritorio (modo oscuro)'],
        ['Flask', '3.0+', 'Framework web (interfaz web)'],
        ['Pandas', '2.0+', 'Manipulación y análisis de datos'],
        ['NumPy', '1.24+', 'Cálculos numéricos y operaciones matriciales'],
        ['SciPy', '1.11+', 'Funciones estadísticas avanzadas'],
        ['Matplotlib', '3.7+', 'Generación de gráficos estadísticos'],
        ['Seaborn', '0.12+', 'Mejora estética de gráficos'],
        ['python-docx', '1.1+', 'Generación de documentos Word en APA 7'],
        ['Bootstrap', '5.3.3', 'Frontend web (interfaz responsiva)'],
        ['Gunicorn', '22.0+', 'Servidor WSGI para despliegue web'],
    ],
    title='Tabla 1',
    note='Las versiones indicadas son las mínimas requeridas. El software fue probado con las versiones más recientes disponibles a julio de 2026.')

add_heading_apa(doc, '3.2 Arquitectura del Sistema', 2)

add_paragraph(doc,
    'El software sigue el patrón de arquitectura Modelo-Vista-Controlador (MVC) en su '
    'versión de escritorio, que separa la lógica de negocio (modelo), la interfaz de '
    'usuario (vista) y la lógica de coordinación (controlador). La versión web sigue '
    'el patrón de rutas-vistas de Flask con plantillas Jinja2. Ambas versiones comparten '
    'los módulos de modelo y exportación, garantizando consistencia en los cálculos '
    'estadísticos y el formato de los reportes.')

make_table(doc,
    ['Directorio/Archivo', 'Función'],
    [
        ['main.py', 'Punto de entrada de la aplicación de escritorio'],
        ['webapp/app.py', 'Fábrica y configuración de la aplicación Flask'],
        ['webapp/routes.py', 'Definición de todas las rutas y endpoints web'],
        ['controller/controllers.py', 'Controladores que coordinan modelo y vista'],
        ['model/sampling.py', 'Cálculos de Muestreo Aleatorio Simple'],
        ['model/statistics.py', 'Clasificación de variables y cálculos estadísticos'],
        ['model/charts.py', 'Generación de gráficos (barras, sectores, histogramas)'],
        ['model/anova.py', 'Módulo de ANOVA unidireccional y bidireccional'],
        ['export/docx_exporter.py', 'Motor de exportación APA 7 a .docx'],
        ['view/main_view.py', 'Ventana principal con panel lateral de navegación'],
        ['view/dataset_view.py', 'Vista de carga y análisis de datasets'],
        ['view/sampling_view.py', 'Vista de la calculadora M.A.S.'],
        ['view/ui_components.py', 'Componentes reutilizables (tablas, tarjetas)'],
        ['webapp/static/css/style.css', 'Estilos CSS con tema oscuro y glassmorphism'],
        ['webapp/templates/', 'Plantillas HTML (Jinja2) para todas las páginas'],
    ],
    title='Tabla 2',
    note='La arquitectura MVC garantiza la separación de responsabilidades y la modularidad del código.')

add_heading_apa(doc, '3.3 Módulos Implementados', 2)

add_heading_apa(doc, '3.3.1 Módulo de Muestreo (M.A.S.)', 3)

add_paragraph(doc,
    'El módulo de Muestreo Aleatorio Simple permite al usuario calcular el tamaño de '
    'muestra óptimo para su investigación. Los parámetros de entrada incluyen el nivel de '
    'confianza (90%, 95% o 99%), la probabilidad de éxito (p), el error admisible (e) y, '
    'opcionalmente, el tamaño de la población (N). Cuando N es proporcionado, el cálculo '
    'incluye la corrección para poblaciones finitas, mostrando tanto la muestra inicial '
    '(n₀) como la muestra corregida (nf). Los resultados se presentan visualmente mediante '
    'tarjetas métricas con los valores de Z, n, n₀ y nf según corresponda.')

add_heading_apa(doc, '3.3.2 Módulo de Procesamiento de Datos', 3)

add_paragraph(doc,
    'Este módulo constituye el núcleo funcional de la aplicación. El flujo de procesamiento '
    'comienza con la carga de un archivo CSV o Excel mediante la zona de arrastrar y soltar '
    'o el explorador de archivos. Una vez cargado el dataset, el sistema ejecuta '
    'automáticamente los siguientes pasos:')

add_paragraph(doc,
    'Clasificación de variables. Cada columna del dataset se analiza para determinar su '
    'tipo: cualitativa nominal, cualitativa ordinal, cuantitativa discreta o cuantitativa '
    'continua. La clasificación se basa en el tipo de dato de Pandas y puede ser '
    'reclasificada manualmente por el usuario.')

add_paragraph(doc,
    'Tablas de frecuencias. Para cada variable cuantitativa continua, se aplica la regla '
    'de Sturges para determinar el número óptimo de intervalos. Para variables cualitativas '
    'y discretas con 15 o menos valores únicos, se construyen tablas no agrupadas. Todas '
    'las tablas incluyen frecuencias absolutas (fi), acumuladas (Fi), relativas (hi), '
    'relativas porcentuales (hi%) y relativas acumuladas porcentuales (Hi%).')

add_paragraph(doc,
    'Cálculo de medidas estadísticas. Para las variables cuantitativas, se calculan las '
    'cinco medidas de tendencia central, las cuatro medidas de dispersión, las medidas de '
    'posición (Q₁-Q₃, D₁, D₅, D₉, P₁₀-P₉₀) y las medidas de forma (asimetría y curtosis). '
    'Para variables cualitativas, se calcula la moda como medida principal.')

add_paragraph(doc,
    'Generación de gráficos. El sistema genera automáticamente gráficos de barras y '
    'sectores para variables cualitativas, gráficos de barras con ojiva para variables '
    'cuantitativas discretas, e histogramas con polígono de frecuencia y ojiva para '
    'variables cuantitativas continuas.')

add_paragraph(doc,
    'Interpretación automática. El sistema genera texto interpretativo en español para '
    'cada variable analizada, discutiendo la simetría, homogeneidad y curtosis de la '
    'distribución, facilitando la redacción del informe final.')

add_heading_apa(doc, '3.3.3 Módulo de Exportación APA 7', 3)

add_paragraph(doc,
    'El motor de exportación genera documentos .docx con formato APA 7 completo. Las '
    'principales características incluyen: tablas sin bordes verticales con bordes '
    'horizontales únicamente en la cabecera y al pie; títulos de tabla numerados y en '
    'cursiva (Tabla 1, Tabla 2, ...); figuras numeradas con títulos en cursiva (Figura 1, '
    'Figura 2, ...); notas explicativas con la palabra "Nota." en cursiva seguida de texto '
    'regular; y formato general con tipografía Times New Roman 12 puntos, márgenes de '
    '2.54 cm e interlineado 1.5.')

add_paragraph(doc,
    'El sistema incluye un contador de exportaciones gratuitas (3 exportaciones), tras '
    'el cual se muestra un sistema de pago integrado mediante Yape (billetera móvil '
    'peruana) por un valor de S/0.30, con verificación de código de operación.')

add_heading_apa(doc, '3.3.4 Módulo de Calificación y Muro de Reseñas', 3)

add_paragraph(doc,
    'El módulo de calificación permite a los usuarios valorar la aplicación mediante un '
    'sistema de 5 estrellas interactivas y un campo de comentarios. Las reseñas se '
    'almacenan en un archivo JSON local (feedback_data.json) con la estructura '
    '{rating, comment, date}. La interfaz incluye un muro de reseñas anónimas que '
    'muestra todas las valoraciones en orden cronológico inverso, con tarjetas de diseño '
    'glassmorphism que exhiben las estrellas, la fecha y el comentario. No se solicita '
    'ni almacena información de nombre o usuario, garantizando el anonimato.')

add_heading_apa(doc, '3.3.5 Módulo de Internacionalización (i18n)', 3)

add_paragraph(doc,
    'El sistema cuenta con un módulo de internacionalización que permite cambiar el idioma '
    'de la interfaz entre español, inglés y portugués. Las traducciones se almacenan en '
    'archivos JSON y se cargan dinámicamente según la selección del usuario. La '
    'arquitectura del módulo permite añadir nuevos idiomas sin modificar el código fuente.')

add_heading_apa(doc, '3.3.6 Versión Web (Flask)', 3)

add_paragraph(doc,
    'Además de la versión de escritorio, se desarrolló una versión web completa utilizando '
    'el framework Flask. Esta versión implementa todas las funcionalidades de análisis '
    'estadístico a través de una API REST y plantillas HTML con Bootstrap 5. La interfaz '
    'web presenta un diseño responsivo con tema oscuro y glassmorphism, panel lateral de '
    'navegación, y menú móvil deslizable. La versión web está preparada para despliegue '
    'en plataformas como HuggingFace Spaces y Heroku mediante contenedores Docker.')

add_heading_apa(doc, '3.4 Interfaz de Usuario', 2)

add_paragraph(doc,
    'La interfaz gráfica fue diseñada siguiendo principios de usabilidad y diseño moderno. '
    'El panel lateral izquierdo proporciona navegación entre las secciones principales: '
    'Inicio (carga de datos), Procesar Dataset, M.A.S., Califícanos, Acerca de y Política '
    'de Privacidad. La zona de trabajo central aloja las vistas correspondientes a cada '
    'sección.')

add_paragraph(doc,
    'Se implementó un tema oscuro profesional con colores corporativos azules (--accent: '
    '#818cf8) y tarjetas con bordes redondeados y efectos glassmorphism (fondo '
    'semitransparente con desenfoque). La zona de carga de archivos utiliza un área de '
    'arrastrar y soltar con realce visual al pasar el cursor. Se incluyen indicadores '
    'de progreso durante el procesamiento de datos para mejorar la experiencia del '
    'usuario.')

add_paragraph(doc,
    'La versión web es totalmente responsiva, con un menú tipo hamburguesa en dispositivos '
    'móviles que se despliega desde la izquierda. El diseño se adapta a diferentes '
    'tamaños de pantalla manteniendo la legibilidad y usabilidad.')

doc.add_page_break()

# ================== CAPÍTULO IV ==================
add_heading_apa(doc, 'Capítulo IV: Resultados y Discusión', 1)

add_paragraph(doc,
    'Para validar el funcionamiento del software, se realizaron pruebas con dos datasets '
    'de ejemplo: dataset_ejemplo_1.csv (40 registros académicos con 7 variables) y '
    'dataset_ejemplo_2.csv (50 registros médicos con 6 variables). En todos los casos, '
    'el sistema clasificó correctamente las variables, construyó las tablas de frecuencias '
    'aplicando la regla de Sturges y calculó las medidas estadísticas esperadas.')

# Table 3
make_table(doc,
    ['Variable', 'Tipo Asignado', 'Fundamento'],
    [
        ['Estudiante', 'Cualitativa Ordinal', 'Identificador único con orden implícito'],
        ['Edad', 'Cuantitativa Discreta', 'Valores enteros con pocas repeticiones'],
        ['Sexo', 'Cualitativa Nominal', 'Categorías sin orden (Masculino/Femenino)'],
        ['Carrera', 'Cualitativa Nominal', 'Categorías sin orden predefinido'],
        ['Nota_Final', 'Cuantitativa Continua', 'Valores decimales continuos (0-20)'],
        ['Horas_Estudio', 'Cuantitativa Continua', 'Valores decimales continuos'],
        ['Asistencias', 'Cuantitativa Continua', 'Valores decimales continuos (0-100%)'],
    ],
    title='Tabla 3',
    note='La clasificación se realiza automáticamente al cargar el dataset, con opción de reclasificación manual.')

add_paragraph(doc,
    'En cuanto al módulo de muestreo M.A.S., se verificaron los cálculos para diferentes '
    'escenarios. Para un nivel de confianza del 95% (Z = 1.96), p = 0.5, e = 0.05 y '
    'población desconocida, el sistema calculó correctamente n = 385. Para una población '
    'conocida de N = 500, aplicó la corrección obteniendo n₀ = 384 y nf = 218. Los '
    'resultados coinciden con los valores teóricos esperados según la fórmula de Cochran '
    '(1977).')

add_paragraph(doc,
    'La generación de gráficos se probó con variables de los dos datasets de ejemplo. '
    'Para variables cualitativas con hasta 10 categorías, se generaron gráficos de barras '
    'y sectores completos. Para variables con más de 10 categorías, el sistema aplicó '
    'correctamente la agrupación en la categoría "Otros". Los histogramas para variables '
    'continuas utilizaron los intervalos calculados por la regla de Sturges, mostrando la '
    'forma de la distribución de los datos.')

add_paragraph(doc,
    'El motor de exportación APA 7 produjo documentos .docx correctamente formateados. '
    'Se verificó que las tablas carecen de bordes verticales, los títulos están en cursiva '
    'numerada, las notas incluyen la palabra "Nota." en cursiva, y el formato general '
    'respeta la tipografía Times New Roman 12 puntos con márgenes de 2.54 cm.')

add_paragraph(doc,
    'En términos de rendimiento, el software procesó un dataset de 40 registros con 7 '
    'variables en menos de 2 segundos, incluyendo la clasificación, los cálculos '
    'estadísticos y la generación de gráficos. La exportación del reporte completo tomó '
    'aproximadamente 3 segundos adicionales. Estos tiempos son aceptables para las '
    'necesidades típicas de estudiantes e investigadores de pregrado.')

add_paragraph(doc,
    'El módulo de calificación y muro de reseñas fue probado exitosamente, permitiendo '
    'el envío y visualización instantánea de reseñas anónimas sin necesidad de recargar '
    'la página. Las pruebas confirmaron que el sistema de estrellas responde correctamente '
    'a los eventos de mouseenter, mouseleave y click, y que los datos se persisten '
    'correctamente en el archivo JSON.')

add_paragraph(doc,
    'La versión web del software fue desplegada exitosamente en HuggingFace Spaces '
    'mediante Docker, demostrando su portabilidad y facilidad de despliegue. La interfaz '
    'web se comportó correctamente en navegadores modernos (Chrome, Firefox, Edge) y en '
    'dispositivos móviles con diferentes tamaños de pantalla.')

doc.add_page_break()

# ================== CONCLUSIONES ==================
add_heading_apa(doc, 'Conclusiones', 1)

add_paragraph(doc,
    '1. Se desarrolló exitosamente el Auto-Analizador de Encuestas, un software estadístico '
    'que automatiza el procesamiento de datos y la generación de reportes en formato APA 7, '
    'cumpliendo con todos los objetivos planteados en el proyecto.')

add_paragraph(doc,
    '2. La implementación de la regla de Sturges para la construcción automática de tablas '
    'de distribución de frecuencias demostró ser precisa y eficiente, generando intervalos '
    'apropiados para la visualización de la distribución de los datos cuantitativos continuos.')

add_paragraph(doc,
    '3. La arquitectura MVC adoptada en la versión de escritorio facilitó la organización '
    'del código y permitió una clara separación entre la lógica de negocio, la interfaz de '
    'usuario y la coordinación del sistema, lo que favorece el mantenimiento futuro y la '
    'extensibilidad del software.')

add_paragraph(doc,
    '4. El motor de exportación APA 7 produce documentos profesionales que cumplen con los '
    'estándares de la séptima edición del manual de publicaciones, eliminando la necesidad '
    'de ajustes manuales de formato y reduciendo significativamente el tiempo de elaboración '
    'de informes académicos.')

add_paragraph(doc,
    '5. Las pruebas realizadas con datasets de ejemplo demostraron la precisión de los '
    'cálculos estadísticos y la robustez del sistema ante diferentes tipos de datos y '
    'configuraciones de entrada.')

add_paragraph(doc,
    '6. La aplicación de la metodología de ingeniería de requisitos, incluyendo la '
    'elicitación mediante entrevistas y observación, la especificación detallada de '
    'requisitos funcionales y no funcionales, y la validación mediante prototipado y '
    'pruebas de usabilidad, garantizó que el producto final se alinee con las necesidades '
    'reales de los estudiantes e investigadores de la UNAS.')

add_paragraph(doc,
    '7. La versión web del software, desarrollada con Flask y Bootstrap, demostró la '
    'viabilidad de ofrecer el análisis estadístico como servicio web, con una interfaz '
    'responsiva y accesible desde cualquier dispositivo.')

doc.add_page_break()

# ================== RECOMENDACIONES ==================
add_heading_apa(doc, 'Recomendaciones', 1)

add_paragraph(doc,
    '1. Se recomienda ampliar el software con módulos adicionales de estadística '
    'inferencial, como pruebas de hipótesis (t de Student, chi-cuadrado) y análisis de '
    'regresión, para cubrir un espectro más amplio de necesidades analíticas.')

add_paragraph(doc,
    '2. Implementar una funcionalidad de importación directa desde Google Sheets y bases '
    'de datos SQL en la versión web, facilitando la integración con sistemas de información '
    'existentes en la universidad.')

add_paragraph(doc,
    '3. Desarrollar un sistema de autenticación de usuarios para la versión web que permita '
    'guardar el historial de análisis y exportaciones realizadas por cada usuario.')

add_paragraph(doc,
    '4. Agregar soporte para más idiomas adicionales, como francés y portugués de Brasil, '
    'para facilitar la colaboración internacional y la publicación en revistas de otros '
    'países.')

add_paragraph(doc,
    '5. Incluir un tutorial interactivo o visitas guiadas dentro de la aplicación web para '
    'facilitar la curva de aprendizaje de nuevos usuarios, especialmente aquellos sin '
    'experiencia previa en software estadístico.')

add_paragraph(doc,
    '6. Implementar pruebas automatizadas de regresión y cobertura de código para mantener '
    'la calidad del software a medida que se añaden nuevas funcionalidades.')

doc.add_page_break()

# ================== REFERENCIAS ==================
add_heading_apa(doc, 'Referencias', 1)

references = [
    'American Psychological Association. (2019). Publication manual of the American Psychological Association (7th ed.). https://doi.org/10.1037/0000165-000',

    'Anderson, D. R., Sweeney, D. J., Williams, T. A., Camm, J. D., & Cochran, J. J. (2019). Estadística para administración y economía (13a ed.). Cengage Learning.',

    'Cochran, W. G. (1977). Sampling techniques (3rd ed.). John Wiley & Sons.',

    'Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering, 9(3), 90-95. https://doi.org/10.1109/MCSE.2007.55',

    'IEEE. (2014). IEEE guide for developing system requirements specifications (IEEE Std 830-2014). IEEE.',

    'ISO/IEC. (2023). ISO/IEC 25010:2023 — Systems and software engineering — Systems and software Quality Requirements and Evaluation (SQuaRE) — Product quality model. ISO.',

    'Lind, D. A., Marchal, W. G., & Wathen, S. A. (2019). Estadística aplicada a los negocios y la economía (17a ed.). McGraw-Hill.',

    'McKinney, W. (2010). Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 51-56.',

    'Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference, 92-96.',

    'Sturges, H. A. (1926). The choice of a class interval. Journal of the American Statistical Association, 21(153), 65-66. https://doi.org/10.1080/01621459.1926.10502161',

    'Van Rossum, G., & Drake, F. L. (2009). Python 3 reference manual. CreateSpace.',

    'Waskom, M. L. (2021). Seaborn: Statistical data visualization. Journal of Open Source Software, 6(60), 3021. https://doi.org/10.21105/joss.03021',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ================== ANEXOS ==================
add_heading_apa(doc, 'Anexos', 1)

add_heading_apa(doc, 'Anexo A: Datasets de Ejemplo', 2)

add_paragraph(doc,
    'Para las pruebas del software se utilizaron dos datasets. El primero contiene 40 '
    'registros académicos de estudiantes universitarios con 7 variables: Estudiante '
    '(identificador), Edad, Sexo, Carrera, Nota_Final, Horas_Estudio y Asistencias. '
    'El segundo contiene 50 registros médicos con 6 variables: Paciente, Edad, Sexo, '
    'Diagnóstico, Días_Hospitalizado y Costo_Total.')

add_paragraph(doc,
    'El software cargó y procesó ambos datasets correctamente, clasificando cada variable '
    'según su tipo y generando las tablas de frecuencias, medidas estadísticas y gráficos '
    'correspondientes.')

add_heading_apa(doc, 'Anexo B: Requisitos Funcionales del Sistema', 2)

add_paragraph(doc,
    'A continuación se presentan los requisitos funcionales identificados durante el '
    'proceso de ingeniería de requisitos para el Auto-Analizador de Encuestas.')

make_table(doc,
    ['ID', 'Nombre', 'Descripción', 'Prioridad'],
    [
        ['RF-01', 'Cargar archivo', 'El sistema debe permitir cargar archivos CSV y Excel mediante arrastrar y soltar o explorador de archivos.', 'Alta'],
        ['RF-02', 'Clasificar variables', 'El sistema debe clasificar automáticamente cada variable del dataset en cualitativa nominal, ordinal, cuantitativa discreta o continua.', 'Alta'],
        ['RF-03', 'Reclasificar variable', 'El sistema debe permitir al usuario cambiar el tipo de una variable mediante un menú desplegable.', 'Alta'],
        ['RF-04', 'Calcular frecuencias', 'El sistema debe construir tablas de distribución de frecuencias aplicando la regla de Sturges para variables continuas.', 'Alta'],
        ['RF-05', 'Calcular medidas', 'El sistema debe calcular medidas de tendencia central, dispersión, posición y forma para cada variable cuantitativa.', 'Alta'],
        ['RF-06', 'Generar gráficos', 'El sistema debe generar gráficos de barras, sectores, histogramas y polígonos de frecuencia.', 'Alta'],
        ['RF-07', 'Interpretar resultados', 'El sistema debe generar texto interpretativo en español para cada variable analizada.', 'Media'],
        ['RF-08', 'Exportar APA 7', 'El sistema debe exportar reportes en formato .docx con formato APA 7.', 'Alta'],
        ['RF-09', 'Calcular muestra', 'El sistema debe calcular el tamaño de muestra mediante M.A.S. para poblaciones conocidas y desconocidas.', 'Alta'],
        ['RF-10', 'Calificar app', 'El sistema debe permitir a los usuarios calificar la aplicación con 1-5 estrellas y dejar comentarios.', 'Media'],
        ['RF-11', 'Ver reseñas', 'El sistema debe mostrar un muro de reseñas anónimas con estrellas, fecha y comentario.', 'Media'],
        ['RF-12', 'Cambiar idioma', 'El sistema debe permitir cambiar entre español, inglés y portugués.', 'Media'],
        ['RF-13', 'Ver resumen', 'El sistema debe mostrar un resumen estadístico de todas las variables cuantitativas.', 'Media'],
        ['RF-14', 'Pagar exportación', 'El sistema debe gestionar el pago mediante Yape tras 3 exportaciones gratuitas.', 'Baja'],
    ],
    title='Tabla A1',
    note='Requisitos funcionales elicitados mediante entrevistas y observación a estudiantes e investigadores de la UNAS.')

doc.add_paragraph()

add_heading_apa(doc, 'Anexo C: Requisitos No Funcionales del Sistema', 2)

make_table(doc,
    ['ID', 'Descripción', 'Característica ISO 25010'],
    [
        ['RNF-01', 'La interfaz debe cargarse en menos de 3 segundos en conexiones de 10 Mbps.', 'Performance efficiency'],
        ['RNF-02', 'El sistema debe procesar datasets de hasta 10,000 registros sin degradación notable.', 'Scalability'],
        ['RNF-03', 'Los datos cargados solo persisten en la sesión del usuario (no se almacenan en BD).', 'Confidentiality'],
        ['RNF-04', 'La interfaz debe ser responsiva y funcional en dispositivos móviles y de escritorio.', 'Inclusivity'],
        ['RNF-05', 'Los mensajes de error deben ser claros y en español, explicando la causa y la acción sugerida.', 'Self-descriptiveness'],
        ['RNF-06', 'Un usuario nuevo debe poder cargar un archivo y ver resultados en menos de 1 minuto sin capacitación.', 'Operability'],
        ['RNF-07', 'El sistema debe estar disponible 24/7 en su versión web desplegada.', 'Availability'],
        ['RNF-08', 'Las reseñas de calificación deben ser anónimas; no se debe solicitar ni almacenar nombre de usuario.', 'Privacy'],
        ['RNF-09', 'El código debe mantener una cobertura de pruebas unitarias superior al 60%.', 'Testability'],
        ['RNF-10', 'El sistema debe soportar los idiomas español, inglés y portugués.', 'Adaptability'],
    ],
    title='Tabla A2',
    note='Requisitos no funcionales basados en el modelo de calidad ISO/IEC 25010:2023.')

add_heading_apa(doc, 'Anexo D: Glosario de Términos', 2)

make_table(doc,
    ['Término', 'Definición'],
    [
        ['Amplitud de intervalo (C)', 'Diferencia entre el límite superior e inferior de un intervalo de clase en una tabla de frecuencias. Se calcula como C = R / m.'],
        ['Asimetría', 'Medida de la forma de una distribución que indica si los datos se extienden más hacia la derecha (positiva) o hacia la izquierda (negativa).'],
        ['Coeficiente de variación (CV)', 'Medida de dispersión relativa que expresa la desviación estándar como porcentaje de la media.'],
        ['Curtosis', 'Medida del apuntamiento de una distribución: leptocúrtica, mesocúrtica o platicúrtica.'],
        ['Decil', 'Medida de posición que divide los datos ordenados en diez partes iguales (D₁ a D₉).'],
        ['Desviación estándar (S)', 'Raíz cuadrada de la varianza. Mide la dispersión promedio de los datos respecto a la media.'],
        ['Estadística descriptiva', 'Rama de la estadística que organiza, resume y presenta datos mediante tablas, gráficos y medidas numéricas.'],
        ['Frecuencia absoluta (fi)', 'Número de veces que aparece un valor o intervalo en el conjunto de datos.'],
        ['Frecuencia acumulada (Fi)', 'Suma de las frecuencias absolutas desde el primer intervalo hasta el intervalo actual.'],
        ['Frecuencia relativa (hi)', 'Proporción de la frecuencia absoluta respecto al total de datos (hi = fi / n).'],
        ['M.A.S.', 'Muestreo Aleatorio Simple. Técnica de muestreo probabilístico donde cada elemento tiene igual probabilidad de ser seleccionado.'],
        ['Media aritmética (X̄)', 'Promedio de todos los valores. Se calcula sumando los valores y dividiendo entre el número de observaciones.'],
        ['Mediana (Me)', 'Valor que divide los datos ordenados en dos mitades iguales (percentil 50).'],
        ['Moda (Mo)', 'Valor que aparece con mayor frecuencia en un conjunto de datos.'],
        ['Percentil', 'Medida de posición que divide los datos ordenados en cien partes iguales (P₁ a P₉₉).'],
        ['Regla de Sturges', 'Fórmula para determinar el número óptimo de intervalos: m = 1 + 3.322 · log₁₀(n).'],
        ['Varianza (S²)', 'Promedio de las desviaciones cuadráticas respecto a la media.'],
        ['Variable cualitativa', 'Variable que expresa categorías o atributos no numéricos.'],
        ['Variable cuantitativa', 'Variable que expresa cantidades numéricas (discreta o continua).'],
        ['Ingeniería de requisitos', 'Proceso sistemático de desarrollar requisitos de software mediante elicitación, especificación, modelado y validación.'],
        ['Elicitación', 'Proceso de descubrir y obtener requisitos a partir de los stakeholders mediante entrevistas, observación y análisis de documentos.'],
    ],
    title='Tabla D1',
    note='Definiciones adaptadas al contexto del software desarrollado.')

# ---- NÚMERO DE PÁGINA ----
add_page_number(doc)

# ======= SAVE =======
doc.save(OUTPUT_PATH)
print(f'Documento generado: {OUTPUT_PATH}')
